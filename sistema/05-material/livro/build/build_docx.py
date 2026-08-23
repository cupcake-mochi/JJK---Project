#!/usr/bin/env python3
"""
Projeto - M · Manual da Guilda
Markdown -> .docx de revisão: só o texto, formatação mínima (título,
negrito, tabela simples), sem nenhuma diagramação de PDF.
"""
import os
import re

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)
MANUAL_DIR = os.path.join(ROOT, "manual")
OUT_DOCX = os.path.join(ROOT, "Projeto-M-Manual-da-Guilda-REVISAO.docx")

CHAPTERS = [
    ("05-introducao.md", "Bem-vindo à Guilda"),
    ("07-glossario.md", "O vocabulário do sistema"),
    ("08-inicio-rapido.md", "Antes da primeira sessão"),
    ("10-como-jogar.md", "Como Jogar"),
    ("11-o-turno.md", "O Turno"),
    ("12-pericias-e-oficios.md", "Perícias e Ofícios"),
    ("15-dano-e-condicoes.md", "Dano, Condições e Cobertura"),
    ("70-descanso-e-recuperacao.md", "Descanso e Recuperação"),
    ("20-criacao-de-personagem.md", "Criação de Personagem"),
    ("25-origens.md", "Origens e Legados"),  # o .md já traz este título
    ("35-caminhos-e-trilhas.md", "Caminhos e Trilhas"),
    ("40-fundamento.md", "Fundamento"),
    ("42-tecnica-marcial.md", "Técnica Marcial"),
    ("45-aptidoes-e-refino.md", "Aptidões e Refino"),
    ("47-bencaos-e-lapidacao.md", "Bênçãos e Lapidação"),
    ("50-equipamento.md", "Equipamento"),
    ("55-ferramenta-amaldicoada.md", "Ferramenta Amaldiçoada"),
    ("60-invocacoes.md", "Invocações"),
    ("65-pactos.md", "Pactos"),
    ("80-experiencia-e-progressao.md", "Experiência e Progressão"),
    ("90-apendice-bloquear.md", "Apêndice · Bloquear"),
]

NOTES_HEADING = "## Notas de revisão"


def split_notes(md_text):
    idx = md_text.find(NOTES_HEADING)
    if idx == -1:
        return md_text, None
    return md_text[:idx].rstrip() + "\n", md_text[idx + len(NOTES_HEADING):].strip()


def add_runs(paragraph, node):
    """Percorre os filhos de um elemento HTML e escreve runs com negrito,
    itálico e monoespaçado onde o markdown original pedia."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                paragraph.add_run(text)
            continue
        name = child.name
        if name in ("strong", "b"):
            r = paragraph.add_run(child.get_text())
            r.bold = True
        elif name in ("em", "i"):
            r = paragraph.add_run(child.get_text())
            r.italic = True
        elif name == "code":
            r = paragraph.add_run(child.get_text())
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif name == "br":
            paragraph.add_run().add_break()
        elif name == "a":
            add_runs(paragraph, child)
        else:
            add_runs(paragraph, child)


HEADING_STYLE = {"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3",
                 "h4": "Heading 4", "h5": "Heading 5"}


def walk(doc, soup):
    for el in soup.find_all(recursive=False):
        name = el.name

        if name in HEADING_STYLE:
            p = doc.add_paragraph(style=HEADING_STYLE[name])
            add_runs(p, el)

        elif name == "p":
            p = doc.add_paragraph()
            add_runs(p, el)

        elif name == "blockquote":
            # citação do markdown = regra em destaque. Aqui vira parágrafo
            # recuado com uma borda esquerda leve, sem cor de fundo.
            for inner in el.find_all(["p"], recursive=False):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                add_runs(p, inner)
            # blockquote sem <p> filho (markdown "extra" às vezes gera assim)
            if not el.find(["p"], recursive=False):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                add_runs(p, el)

        elif name in ("ul", "ol"):
            style = "List Bullet" if name == "ul" else "List Number"
            for li in el.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style)
                add_runs(p, li)

        elif name == "table":
            rows = el.find_all("tr")
            if not rows:
                continue
            ncols = max(len(r.find_all(["th", "td"])) for r in rows)
            t = doc.add_table(rows=0, cols=ncols)
            t.style = "Light Grid Accent 1"
            for r in rows:
                cells = r.find_all(["th", "td"])
                row = t.add_row()
                is_header = bool(r.find("th"))
                for i in range(ncols):
                    cell = row.cells[i]
                    cell.paragraphs[0].clear() if cell.paragraphs[0].runs else None
                    p = cell.paragraphs[0]
                    if i < len(cells):
                        add_runs(p, cells[i])
                    if is_header:
                        for run in p.runs:
                            run.bold = True
            doc.add_paragraph()

        elif name == "hr":
            doc.add_paragraph("* * *").alignment = WD_ALIGN_PARAGRAPH.CENTER

        else:
            # qualquer outra coisa: percorre por dentro, sem criar bloco novo
            walk(doc, el)


def md_to_soup(md_text):
    html = markdown.markdown(
        md_text, extensions=["tables", "extra", "sane_lists", "attr_list"]
    )
    return BeautifulSoup(html, "html.parser")


def main():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 20), (2, 16), (3, 13), (4, 11), (5, 11)]:
        st = doc.styles[f"Heading {level}"]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.italic = False

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Projeto M — Manual da Guilda")
    r.bold = True
    r.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Documento de revisão de texto — sem formatação final, só o conteúdo")
    r.italic = True

    doc.add_page_break()

    notas_geral = []

    for filename, titulo in CHAPTERS:
        path = os.path.join(MANUAL_DIR, filename)
        if not os.path.exists(path):
            print(f"  AVISO: {filename} não encontrado, pulando.")
            continue
        with open(path, encoding="utf-8") as f:
            raw = f.read()
        corpo_md, notas_md = split_notes(raw)
        # tira o "# Título" do topo do markdown-fonte: o título já vira Heading 1 aqui
        corpo_md = re.sub(r"^#\s+.*\n", "", corpo_md, count=1)

        h = doc.add_paragraph(style="Heading 1")
        h.add_run(titulo)

        soup = md_to_soup(corpo_md)
        walk(doc, soup)

        if notas_md:
            notas_geral.append((titulo, notas_md))

    doc.save(OUT_DOCX)
    print(f"docx: {OUT_DOCX}")
    print(f"  {len(CHAPTERS)} capítulos processados; {len(notas_geral)} tinham notas de revisão (não incluídas).")


if __name__ == "__main__":
    main()
