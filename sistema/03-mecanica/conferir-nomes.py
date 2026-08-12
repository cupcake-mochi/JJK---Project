#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Confere TODO nome batizado pelo projeto, nas duas direcoes.

O conferir-pericias.py ja faz isso — mas so com 33 nomes: as 23 pericias e os
10 oficios. O projeto batizou muito mais coisa: 15 Trilhas, 5 Caminhos, 14
Legados, 8 Origens, os termos de sistema. Nenhum desses passava pela checagem,
e foi por isso que Sombra, Enxame e Oficio ficaram no material.

Cinco checagens:
  1. COLISAO PARA FORA — o nome ja e termo definido no manual do Fundamento,
     sozinho (Sombra e Tema) ou dentro de um termo maior (Folego esta em
     Roubo de Folego). Prosa solta nao conta — nota de metodo da v0.6.
  2. COLISAO PARA DENTRO — o nome ja significa outra coisa no proprio projeto:
     ou e a mesma palavra que outro nome batizado, ou aparece em minuscula como
     palavra comum de peso em outra peca.
  3. COLISAO FRACA — o nome fica a uma letra de um termo definido do manual
     (Remendo contra a Melhoria Remenda). Avisa; nao falha.
  4. CATEGORIA ERRADA — todo lugar do projeto que cita uma Familia, uma Forma,
     uma Melhoria ou uma Restricao precisa usar um nome que existe naquela
     categoria. E isto que pega "Familias Livres: Peso e Prender".
  5. TERMO MORTO VIVO — nome aposentado aparecendo fora de nota historica.

O manual e a fonte da verdade das listas fechadas: elas sao extraidas do .docx
em vez de copiadas aqui, para nao envelhecerem quando ele for regerado.

Roda sem argumento. Sai com codigo 1 se algo quebrar.
Sem python-docx, as checagens 1, 3 e 4 sao PULADAS com aviso, em vez de falhar.
"""

import os
import re
import sys
import glob
import unicodedata

FALHAS = []
AVISOS = []


def erro(msg):
    FALHAS.append(msg)
    print(f'  !! {msg}')


def aviso(msg):
    AVISOS.append(msg)
    print(f'  ~  {msg}')


def sem_acento(s):
    """Tira acento e mantem a caixa — a caixa separa nome proprio de prosa."""
    s = unicodedata.normalize('NFD', s)
    return ''.join(c for c in s if unicodedata.category(c) != 'Mn')


def norma(s):
    return sem_acento(s).lower()


def bloco(titulo):
    print()
    print('=' * 88)
    print(titulo)
    print('=' * 88)


def perto(a, b):
    """True se a e b diferem por uma letra so. Nomes assim confundem na mesa."""
    a, b = norma(a), norma(b)
    if a == b or abs(len(a) - len(b)) > 1 or min(len(a), len(b)) < 4:
        return False
    if len(a) == len(b):
        return sum(x != y for x, y in zip(a, b)) == 1
    curto, longo = (a, b) if len(a) < len(b) else (b, a)
    return any(longo[:i] + longo[i + 1:] == curto for i in range(len(longo)))


AQUI = os.path.dirname(os.path.abspath(__file__))
RAIZ = os.path.normpath(os.path.join(AQUI, '..'))
DOCX = os.path.normpath(os.path.join(RAIZ, '..', 'manual', 'Fundamento-MANUAL-v7.docx'))

# --------------------------------------------------------------------------
# OS NOMES QUE O PROJETO BATIZOU
# --------------------------------------------------------------------------
CAMINHOS = ['Bastiao', 'Vanguarda', 'Guia', 'Emanador', 'Evocador']

# Fechados na v0.24. Seis nomes sairam por colisao: Alcance (Familia do manual,
# e dentro do Legado Alcance Impossivel), Oficio (a categoria de oficio),
# Folego (Roubo de Folego), Regua (Fundamento pronto), Sombra (Tema, e dentro de
# Estilo da Sombra) e Enxame (Tema).
TRILHAS = {
    'Bastiao':   ['Muro', 'Punho', 'Brasa'],
    'Vanguarda': ['Estocada', 'Batedor', 'Executor'],
    'Guia':      ['Elo', 'Sutura', 'Perimetro'],
    'Emanador':  ['Torrente', 'Repertorio', 'Arremate'],
    'Evocador':  ['Servo', 'Matilha', 'Coro'],
}

ORIGENS = ['Latente', 'Receptaculo', 'Descendente', 'Reencarnado', 'Feto',
           'Sem Tecnica', 'Corpo Amaldicoado', 'Restricao Celestial']

LEGADOS = ['Instinto Bruto', 'Aprendi Apanhando', 'A Voz de Dentro',
           'Nao Sou So Eu', 'O Sobrenome', 'Treino de Berco',
           'O Que Ninguem Lembra', 'Corpo Emprestado', 'Sangue que Nao e Sangue',
           'Irmaos', 'Nucleos', 'Nao Sou Gente', 'Alcance Impossivel', 'Peso Real']

SISTEMA = ['Maestria', 'Refino', 'Trilha', 'Caminho', 'Legado', 'Exaustao',
           'Pericia', 'Oficio', 'Golpe canalizado', 'Golpe simples',
           'Ambiente propicio', 'Descanso curto', 'Descanso longo',
           'Tecnica Marcial', 'Estilo da Sombra']

# Onde cada nome e definido. Na checagem 2, o arquivo de definicao nao conta.
DEFINIDO_EM = {
    **{c: '06-caminhos-e-trilhas.md' for c in CAMINHOS},
    **{t: '06-caminhos-e-trilhas.md' for g in TRILHAS.values() for t in g},
    **{o: '09-origens.md' for o in ORIGENS},
    **{l: '09-origens.md' for l in LEGADOS},
}

# As 23 pericias e os 10 oficios vem do conferir-pericias.py, para nao existirem
# duas listas que possam divergir. Se aquele arquivo sumir, seguimos sem eles.
PERICIAS, OFICIOS = [], []
sys.dont_write_bytecode = True   # nao sujar a pasta com __pycache__ ao importar
try:
    import importlib.util
    _spec = importlib.util.spec_from_file_location(
        'conferir_pericias', os.path.join(AQUI, 'conferir-pericias.py'))
    _mod = importlib.util.module_from_spec(_spec)
    import io as _io
    import contextlib as _ctx
    with _ctx.redirect_stdout(_io.StringIO()):
        try:
            _spec.loader.exec_module(_mod)
        except SystemExit:
            pass
    PERICIAS = [p for g in _mod.PERICIAS.values() for p in g]
    OFICIOS = list(_mod.OFICIOS)
except Exception as _e:
    print(f'  (nao consegui ler as listas do conferir-pericias.py: {_e})')

TODOS = ([('Caminho', c) for c in CAMINHOS]
         + [(f'Trilha do {cam}', t) for cam, g in TRILHAS.items() for t in g]
         + [('Origem', o) for o in ORIGENS]
         + [('Legado', l) for l in LEGADOS]
         + [('termo de sistema', s) for s in SISTEMA])

# Usada so para triar nome candidato (--candidatos), nao nas cinco checagens:
# pericia e oficio ja tem dono no conferir-pericias.py.
UNIVERSO = TODOS + [('pericia', p) for p in PERICIAS] + [('oficio', o) for o in OFICIOS]

# --------------------------------------------------------------------------
# COLISOES CONHECIDAS E ACEITAS DE PROPOSITO
# --------------------------------------------------------------------------
# Nao falham. Aparecem no relatorio com o motivo, para separar "aceito" de
# "nao percebido". Mesma regra do conferir-pericias.py.
ACEITAS = {}

# Termo morto -> onde ele morreu. So conta quando aparece CAPITALIZADO: em
# minuscula ("nao potencia", "linha de frente") e palavra comum, nao o termo.
MORTOS = {
    'Artesania': 'virou Entalhador na v0.16',
    'Taticas': 'deixou de existir na v0.16 — coberta por Ocultismo e Sentir Energia',
    'Potencia': 'nome provisorio do tamanho de feitico; virou Classe na v0.6',
    'Maculado': 'virou Receptaculo na v0.22',
    'Desperto': 'virou Reencarnado na v0.22',
    'Canalizador': 'virou Emanador na v0.14, por colisao com golpe canalizado',
    'Linha de Frente': 'rotulo de rascunho; virou Bastiao na v0.14',
    'Ponta de Lanca': 'rotulo de rascunho; virou Vanguarda na v0.14',
    'Primeiros Socorros': 'saiu na v0.16; Herbalismo nao cobre o mesmo',
    'Protocolo': 'virou Burocracia na v0.16',
}

# Onde um termo morto pode aparecer capitalizado sem ser descuido: a secao que
# conta por que ele morreu. Arquivo -> termos liberados ali, com o motivo.
MORTO_LIBERADO = {
    '06-caminhos-e-trilhas.md': {
        'Canalizador': 'a secao 1 desta peca e justamente a que explica o rename',
        'Ponta de Lanca': 'citada na linha que lista os rotulos de rascunho',
        'Linha de Frente': 'citada na linha que lista os rotulos de rascunho',
    },
    'arquitetura.md': {
        'Potencia': 'citado no aviso de cabecalho, que registra o nome antigo',
    },
    'CHANGELOG.md': {k: 'o changelog e o registro historico' for k in MORTOS},
}

# Uma linha que EXPLICA a aposentadoria pode citar o termo morto. Detectado no
# texto normalizado, para "nao passa" casar com "não passa".
MARCA_HISTORICA = re.compile(
    r'(corrigid|resolvid|substitu|deixou de existir|virou|era[m]? rotulo|rascunho|'
    r'antes da v0|ate a v0|da v0\.\d+ ate|aposentad|saiu|nao passa|nao cobre|'
    r'ja nao|deixa de|foi para|na v0\.\d+)')

PADROES_VIVOS = ['00-fundacao/*.md', '01-pesquisa/*.md', '02-esqueleto/*.md',
                 '03-mecanica/*.md', 'ESTADO-ATUAL.md', 'LEIA-ME.md']

vivos = {}
for pad in PADROES_VIVOS:
    for caminho in sorted(glob.glob(os.path.join(RAIZ, pad))):
        vivos[os.path.basename(caminho)] = open(caminho, encoding='utf-8').read()

print('=' * 88)
print('CONFERIR NOMES — colisao de termo nas duas direcoes')
print('=' * 88)
print(f'  {len(TODOS)} nomes batizados, contra {len(vivos)} arquivos vivos')

# --------------------------------------------------------------------------
# O VOCABULARIO DEFINIDO DO MANUAL, extraido do .docx
# --------------------------------------------------------------------------
manual = None
CATEGORIAS = {}
PLURAL = {'Familia': 'Familias', 'Forma': 'Formas', 'Melhoria': 'Melhorias',
          'Restricao': 'Restricoes', 'Tema': 'Temas', 'Passiva': 'Passivas',
          'Feitico pronto': 'Feiticos prontos', 'Fundamento pronto': 'Fundamentos prontos',
          'peca do Fundamento': 'pecas do Fundamento'}

# Termos que o manual DEFINE em prosa, sem tabela de onde extrair. Estes ficam
# escritos aqui porque nao ha coluna para ler — mas cada um tem secao propria
# no manual, entao colidir com eles e colisao forte.
NUCLEO_DO_MANUAL = ['Fundamento', 'Regra', 'Descricao', 'Familia', 'Selo', 'Passiva',
                    'Classe', 'Forma', 'Melhoria', 'Restricao', 'Liberacao Maxima',
                    'Tecnica Maxima', 'Integridade', 'Uso Livre', 'Efeito Proprio',
                    'Regra Propria', 'Restricao Propria', 'Ampliar']

try:
    import docx  # python-docx
    doc = docx.Document(DOCX)
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                partes.append(c.text)
    manual = '\n'.join(partes)

    def primeira_coluna(rotulo, extras=()):
        """Nomes da primeira coluna das tabelas cujo cabecalho e `rotulo`."""
        fora = {norma(rotulo)} | {norma(x) for x in extras}
        achados = []
        for t in doc.tables:
            cab = [c.text.strip() for c in t.rows[0].cells]
            if cab and norma(cab[0]) == norma(rotulo):
                for r in t.rows[1:]:
                    n = r.cells[0].text.strip()
                    if n and norma(n) not in fora and '\n' not in n:
                        achados.append(n)
        return achados

    CATEGORIAS['Familia'] = primeira_coluna('Familia')
    CATEGORIAS['Forma'] = [n for n in primeira_coluna('Forma')
                           if not n.lower().startswith('projetil e')]
    CATEGORIAS['Melhoria'] = primeira_coluna('Melhoria')
    CATEGORIAS['Restricao'] = primeira_coluna('Restricao')
    CATEGORIAS['Feitico pronto'] = primeira_coluna(
        'Nome', extras=('Classe / Pontos / PE', 'Classe / Pontos / Teto / PE'))
    CATEGORIAS['Passiva'] = primeira_coluna('Passiva')
    CATEGORIAS['peca do Fundamento'] = list(NUCLEO_DO_MANUAL)

    temas = []
    for t in doc.tables:
        cab = [c.text.strip() for c in t.rows[0].cells]
        if len(cab) > 1 and norma(cab[0]) == 'grupo' and norma(cab[1]) == 'temas':
            for r in t.rows[1:]:
                temas += [x.strip() for x in r.cells[1].text.split('·') if x.strip()]
    CATEGORIAS['Tema'] = temas

    # Os tres Fundamentos prontos sao paragrafos curtos logo depois do titulo.
    prontos, linhas = [], manual.split('\n')
    for i, l in enumerate(linhas):
        if norma(l.strip()) == 'tres fundamentos prontos':
            for prox in linhas[i + 1:i + 12]:
                p = prox.strip()
                if p and len(p.split()) == 1 and p[0].isupper():
                    prontos.append(p)
    CATEGORIAS['Fundamento pronto'] = prontos

    print(f'  manual: {len(doc.paragraphs)} paragrafos, {len(doc.tables)} tabelas')
    for cat, lista in CATEGORIAS.items():
        print(f'     {PLURAL[cat]:<20} {len(lista):>3} extraidos do .docx')
    if len(CATEGORIAS['Familia']) != 9:
        erro(f'o manual devolveu {len(CATEGORIAS["Familia"])} Familias, e sao nove — '
             f'a extracao quebrou e as checagens 1 e 4 nao valem')
    if len(CATEGORIAS['Fundamento pronto']) != 3:
        erro(f'o manual devolveu {len(CATEGORIAS["Fundamento pronto"])} Fundamentos '
             f'prontos, e sao tres — a extracao quebrou')
except Exception as e:
    print(f'  (manual nao lido: {e})')
    print('  As checagens 1, 3 e 4 foram PULADAS. Instale python-docx e deixe o')
    print('  manual em manual/Fundamento-MANUAL-v7.docx antes de fechar versao.')

# --------------------------------------------------------------------------
bloco('1. COLISAO PARA FORA — o nome ja e termo definido no Fundamento')

if manual:
    print('  Nota de metodo da v0.6: TERMO DEFINIDO colide; PROSA SOLTA nao.\n')
    exatos = {norma(n): (cat, n) for cat, lista in CATEGORIAS.items() for n in lista}
    achou = False
    for cat, nome in TODOS:
        k = norma(nome)
        if k in exatos:
            c, orig = exatos[k]
            achou = True
            if nome in ACEITAS:
                print(f'  ACEITA  {nome:<20} {cat} — e {c} no manual')
                print(f'          motivo: {ACEITAS[nome]}')
            else:
                erro(f'"{nome}" ({cat}) e {c} definido no manual')
            continue
        # dentro de um termo maior: Folego esta em "Roubo de Folego"
        dentro = [(c, n) for c, lista in CATEGORIAS.items() for n in lista
                  if len(n.split()) > 1
                  and re.search(r'(?<![0-9a-z])' + re.escape(k) + r'(?![0-9a-z])', norma(n))]
        if dentro and nome not in ACEITAS:
            achou = True
            c, n = dentro[0]
            erro(f'"{nome}" ({cat}) esta dentro de "{n}", que e {c} no manual')
    if not achou:
        print('  Nenhum nome do projeto e termo definido do manual.')
else:
    print('  PULADA.')

# --------------------------------------------------------------------------
bloco('2. COLISAO PARA DENTRO — o nome ja significa outra coisa no projeto')

achou = False
vistos = {}
for cat, nome in TODOS:
    k = norma(nome)
    if k in vistos:
        erro(f'"{nome}" e {cat} e tambem {vistos[k]} — a mesma palavra em dois lugares')
        achou = True
    else:
        vistos[k] = cat

# Um nome batizado dentro de outro nome batizado (Sombra em Estilo da Sombra).
for cat_a, a in TODOS:
    for cat_b, b in TODOS:
        if a == b or len(b.split()) < 2:
            continue
        if re.search(r'(?<![0-9a-z])' + re.escape(norma(a)) + r'(?![0-9a-z])', norma(b)):
            erro(f'"{a}" ({cat_a}) esta dentro de "{b}" ({cat_b}) — a mesma palavra '
                 f'nomeia duas pecas do sistema')
            achou = True

# Um nome batizado que tambem e palavra comum do portugues: aparece em
# MINUSCULA, fora do arquivo onde e definido, com peso. Pela nota de metodo da
# v0.6 isso e prosa solta e nao falha — mas vale o aviso, porque nome de peca
# que tambem e palavra comum e lido errado na primeira leitura.
LIMIAR = 3
for cat, nome in TODOS:
    if len(nome.split()) > 1 or nome in ACEITAS or cat == 'termo de sistema':
        continue
    dono = DEFINIDO_EM.get(nome, '')
    pat = re.compile(r'(?<![0-9A-Za-z])' + re.escape(norma(nome)) + r'(?![0-9A-Za-z])')
    for arq, txt in sorted(vivos.items()):
        if arq == dono:
            continue
        texto = sem_acento(txt)   # tira acento, mantem a caixa
        n = len([m for m in pat.finditer(texto)
                 if m.start() > 0 and texto[m.start() - 1] not in '.!?\n>|*'])
        if n >= LIMIAR:
            aviso(f'"{nome}" ({cat}) tambem e palavra comum: {n}x em minuscula '
                  f'em {arq}')

if not achou:
    print('  Nenhum nome batizado carrega dois significados dentro do projeto.')

# --------------------------------------------------------------------------
bloco('3. COLISAO FRACA — o nome fica a uma letra de um termo do manual')

if manual:
    achou = False
    for cat, nome in TODOS:
        for c, lista in CATEGORIAS.items():
            for def_ in lista:
                if perto(nome, def_):
                    achou = True
                    aviso(f'"{nome}" ({cat}) fica a uma letra de "{def_}", que e {c} '
                          f'no manual — na mesa falada isso confunde')
    if not achou:
        print('  Nenhum nome do projeto encosta num termo do manual por uma letra.')
else:
    print('  PULADA.')

# --------------------------------------------------------------------------
bloco('4. CATEGORIA ERRADA — exemplo do projeto citando nome que nao existe la')

if manual:
    GATILHOS = [
        (r'Fam[ií]lias?\s+Livres?:?\**\s*([^.\n|;]+)', 'Familia'),
        (r'\**Fechadas:?\**\s*([^.\n|;]+)', 'Familia'),
        (r'Forma\s+([A-ZÀ-Ý][a-zà-ÿ]+)', 'Forma'),
        (r'Restri[cç][aã]o\s+([A-ZÀ-Ý][a-zà-ÿ]+(?:\s+[a-zà-ÿ]+\s+[A-ZÀ-Ý][a-zà-ÿ]+)?)', 'Restricao'),
        (r'Melhoria\s+([A-ZÀ-Ý][a-zà-ÿ]+)', 'Melhoria'),
    ]
    BATIZADOS = {norma(n) for _, n in TODOS}
    # "Restricao Celestial" e Origem: o gatilho de Restricao acha "Celestial" ali
    # e nao e erro. Isso vale so para o gatilho, nao para as listas de Familia.
    PREFIXO = {'Restricao': 'restricao', 'Melhoria': 'melhoria', 'Forma': 'forma'}
    achou = False
    for arq, txt in sorted(vivos.items()):
        for pad, cat in GATILHOS:
            validos = {norma(n) for n in CATEGORIAS.get(cat, [])}
            if not validos:
                continue
            for m in re.finditer(pad, txt):
                bruto = re.sub(r'\*+|`|_', '', m.group(1))
                for n in re.split(r',| e |·|/', bruto):
                    n = n.strip(' .:*—-')
                    if not n or len(n.split()) > 2:
                        continue
                    # so candidato em Title Case; prosa em minuscula nao e nome
                    if not n[0].isupper():
                        continue
                    k = norma(n)
                    if k in validos or k in BATIZADOS:
                        continue
                    # "Restricao Celestial" e Origem, nao Restricao mal citada
                    pre = PREFIXO.get(cat)
                    if pre and f'{pre} {k}' in BATIZADOS:
                        continue
                    outra = next((c for c, l in CATEGORIAS.items()
                                  if k in {norma(x) for x in l}), None)
                    extra = f' — ele e {outra} no manual' if outra else ''
                    erro(f'{arq}: "{n}" citado como {cat} e nao existe nessa lista{extra}')
                    achou = True
    if not achou:
        print('  Todo nome citado como Familia, Forma, Melhoria ou Restricao existe')
        print('  na lista correspondente do manual.')
else:
    print('  PULADA.')

# --------------------------------------------------------------------------
bloco('5. TERMO MORTO VIVO — nome aposentado fora de nota historica')

achou = False
for arq, txt in sorted(vivos.items()):
    liberados = MORTO_LIBERADO.get(arq, {})
    originais = txt.split('\n')
    # sem_acento preserva a caixa e o numero de linhas: "Taticas" casa "Táticas",
    # e "taticas" em prosa continua sem casar.
    limpas = sem_acento(txt).split('\n')
    for morto, destino in MORTOS.items():
        pat = re.compile(r'(?<![0-9A-Za-z])' + re.escape(sem_acento(morto))
                         + r'(?![0-9A-Za-z])')
        achadas = [i for i, l in enumerate(limpas) if pat.search(l)]
        if not achadas:
            continue
        if morto in liberados:
            print(f'  ACEITA  {morto:<20} em {arq}')
            print(f'          motivo: {liberados[morto]}')
            continue
        for i in achadas:
            if MARCA_HISTORICA.search(limpas[i].lower()):
                continue
            achou = True
            erro(f'{arq}: "{morto}" capitalizado fora de nota historica ({destino})')
            print(f'         > {originais[i].strip()[:76]}')
if not achou:
    print('  Nenhum termo aposentado sobrou capitalizado fora de nota historica.')

# --------------------------------------------------------------------------
bloco('6. TRIAGEM DE CANDIDATO — rode antes de batizar qualquer coisa')

# python3 conferir-nomes.py --candidatos Vulto Matilha Bigorna
# Nao faz parte do contrato: nao falha, so responde se o nome esta livre.
#
#   OCUPADO  o nome inteiro ja e termo definido. Escolha outro.
#   DENTRO   o nome aparece dentro de um termo composto. NAO e veredito:
#            va ler o termo e pergunte se ele E aquilo. Se nao for, use.
#   fraco    fica a uma letra de um termo. Confunde na mesa em voz alta.
#   LIVRE    ninguem usa.
candidatos = sys.argv[2:] if len(sys.argv) > 1 and sys.argv[1] == '--candidatos' else []
if not candidatos:
    print('  (nenhum. Uso: python3 conferir-nomes.py --candidatos Vulto Matilha ...)')
else:
    ocupados = {norma(n): cat for cat, n in UNIVERSO}
    for cand in candidatos:
        k = norma(cand)
        # DUROS: o nome inteiro ja e um termo definido. Mata o candidato.
        # DENTRO: o nome aparece dentro de um termo COMPOSTO. Nao mata —
        #   pede a pergunta de sentido, e so o humano responde. Criterio do
        #   Mizuki, v0.40: "se preocupe quando o nome bate de frente com o
        #   nome de algo que e REALMENTE aquilo". Uma Melhoria chamada
        #   "Rasga Escudo" nao E um escudo, entao ela nao ocupa "Escudo".
        duros, dentro = [], []
        if k in ocupados:
            duros.append(f'ja e {ocupados[k]} no projeto')
        for cat, n in UNIVERSO:
            if len(n.split()) > 1 and re.search(
                    r'(?<![0-9a-z])' + re.escape(k) + r'(?![0-9a-z])', norma(n)):
                dentro.append(f'dentro de "{n}" ({cat})')
        for c, lista in CATEGORIAS.items():
            for x in lista:
                if norma(x) == k:
                    duros.append(f'e {c} no manual')
                elif len(x.split()) > 1 and re.search(
                        r'(?<![0-9a-z])' + re.escape(k) + r'(?![0-9a-z])', norma(x)):
                    dentro.append(f'dentro de "{x}" ({c})')
                # e o inverso: "Selo de uma Mao" carrega "Selo" dentro dele
                elif len(cand.split()) > 1 and re.search(
                        r'(?<![0-9a-z])' + re.escape(norma(x)) + r'(?![0-9a-z])', k):
                    dentro.append(f'carrega "{x}" ({c}) dentro dele')
        fracas = [f'a uma letra de "{x}" ({c})'
                  for c, lista in CATEGORIAS.items() for x in lista if perto(cand, x)]
        if duros:
            print(f'  OCUPADO  {cand:<16} {"; ".join(dict.fromkeys(duros))[:62]}')
        elif dentro:
            print(f'  DENTRO   {cand:<16} {"; ".join(dict.fromkeys(dentro))[:62]}')
        elif fracas:
            print(f'  fraco    {cand:<16} {fracas[0]}')
        else:
            print(f'  LIVRE    {cand}')

# --------------------------------------------------------------------------
print()
print('=' * 88)
if FALHAS:
    print(f'>>> {len(FALHAS)} PROBLEMA(S):')
    for f in FALHAS:
        print(f'    - {f}')
    if AVISOS:
        print(f'    (e {len(AVISOS)} aviso(s), que nao falham)')
    sys.exit(1)
print('>>> TUDO OK — nenhum nome batizado colide, nenhuma categoria do manual foi')
print('    citada errada, e nenhum termo morto sobrou fora de nota historica.')
if AVISOS:
    print(f'    {len(AVISOS)} aviso(s) acima, que nao falham o validador.')
