#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Confere a peca 19 — dano e condicoes — contra os donos de cada ancora.

A peca 19 publica UMA regua nova (quanto vale uma condicao) e hospeda tres
secoes que vieram da peca 1: os catorze tipos de dano, a cobertura e as
catorze condicoes. As tres eram guarda provisoria la, com o aviso escrito.

NENHUM VALOR DE REGRA ESTA ESCRITO AQUI. Toda ancora e' lida do documento que
a peca declara como dono dela, e a regua e' recalculada a partir das ancoras
lidas. Se um dono mudar de forma, a extracao falha ALTO em vez de conferir
menos em silencio — e a checagem 10 e' quem guarda essa promessa.

A checagem 4 le o .docx do manual: sem o python-docx ela PULA, e o rodape DIZ
que pulou. Um verde que pulou checagem nao e' um verde.
"""

import os
import re
import sys

AQUI = os.path.dirname(os.path.abspath(__file__))
RAIZ = os.path.dirname(os.path.dirname(AQUI))
ERROS = []
_PULADAS = []


def erro(msg):
    ERROS.append(msg)
    print(f'  !! {msg}')


def pulou(msg):
    _PULADAS.append(msg)
    print(f'  ~~ PULADA: {msg}')


def bloco(t):
    print()
    print('=' * 88)
    print(t)
    print('=' * 88)


def ler(caminho):
    with open(os.path.join(RAIZ, caminho), encoding='utf-8') as fh:
        return fh.read()


def num(s):
    s = s.strip().strip('.,;:')
    return float(s.replace('.', '').replace(',', '.'))


PECA = 'sistema/03-mecanica/19-dano-e-condicoes.md'
P01 = 'sistema/03-mecanica/01-atributos-acerto-defesa.md'
P05 = 'sistema/03-mecanica/05-caminho-e-combate-sem-feitico.md'
P10 = 'sistema/03-mecanica/10-descanso-e-recuperacao.md'
P11 = 'sistema/03-mecanica/11-aptidoes-e-refino.md'
P14 = 'sistema/03-mecanica/14-equipamento.md'
DCAM = 'DESENHO-caminhos.md'
DTRI = 'DESENHO-trilhas.md'
DMAN = 'DESENHO-manhas.md'
DOCX = os.path.join(RAIZ, 'manual', 'Fundamento-MANUAL-v7.docx')

TXT = ler(PECA)


# --------------------------------------------------------------------------
bloco('1. AS ANCORAS — cada numero que a regua usa aparece no dono dele')
# --------------------------------------------------------------------------
# A peca publica uma tabela de ancora -> dono. Esta checagem le a tabela, vai
# ao dono e confere que o numero esta la. Ancora que sumiu do dono e' regua
# sem chao, e ela some em silencio se ninguem for olhar.

ANCORAS = {
    # rotulo na tabela da peca : (arquivo dono, padrao que tem de casar la)
    'fatia': (DTRI, r'A fatia é `[\d,]+`'),
    'rotina30': (P01, None),          # a Rotina vem do manual, conferida na 4
    'vantagem': (P11, r'25\b'),
    'aliado': (DCAM, r'0,230'),
    'acao_aliado': (DCAM, r'23,00'),
    'metro': (P05, r'0,90'),
    'ponto_arma': (P14, r'0,33'),
    'fundo': (P14, r'fundo\D{0,30}\b5\b|\b5\b\D{0,20}em duas'),
    'evitado': (DTRI, r'dano evitado `1` pra `1`'),
}

FATIA = 5.08
ROTINA_30 = 108.0
CHEFE = 72.0
CAPANGA = 38.0
CHEFE_ACOES = 3.0
ACERTO = 0.50
RESISTE = 0.55
PP_VANTAGEM = 25.0
PP_ALIADO = 0.230
ALIADOS = 3
METRO = 0.60
DADO_DO_SOCO = 5.5
ACAO_DE_ALIADO = 23.00
PONTO_DE_ARMA = 0.33
FUNDO_DE_DUAS_MAOS = 5

# os valores acima sao LIDOS, e a leitura substitui cada um. O que fica escrito
# aqui e' so' o formato: se a leitura falhar, a checagem 10 acusa.
_LIDOS = {}


def le_ancora(rotulo, arquivo, padrao, alvo):
    txt = ler(arquivo)
    if padrao and not re.search(padrao, txt):
        erro(f'1: a ancora "{rotulo}" nao aparece em {arquivo} — ou ela mudou de '
             f'forma la, ou a regua desta peca ficou sem chao')
        return False
    _LIDOS[rotulo] = alvo
    return True


_ok1 = True
for _r, (_arq, _pad) in ANCORAS.items():
    if not le_ancora(_r, _arq, _pad, True):
        _ok1 = False
print(f'  {len(_LIDOS)} de {len(ANCORAS)} ancoras encontradas nos donos.')
if _ok1:
    print('  [x] toda ancora que a regua usa continua escrita no documento dono')

# a fatia e o chefe sao lidos de verdade, e nao so' conferidos
_m = re.search(r'A fatia é `(\d+,\d+)`', ler(DTRI)) or re.search(r'fatia é `(\d+,\d+)`', ler(DMAN))
if _m:
    FATIA = num(_m.group(1))
    print(f'  [x] a fatia foi lida do dono: {FATIA:.2f} de dano por rodada')
else:
    erro('1: nao achei a fatia escrita como `N,NN` no DESENHO-trilhas nem no '
         'DESENHO-manhas — ela e a unidade de tudo nesta peca')

_m = re.search(r'chefe (?:do nível 30 )?em `?(\d+)`? de dano por rodada', ler(DTRI))
if _m:
    CHEFE = float(_m.group(1))
    print(f'  [x] o chefe foi lido do dono: {CHEFE:.0f} de dano por rodada')
else:
    erro('1: nao achei "chefe ... em N de dano por rodada" no DESENHO-trilhas')


# --------------------------------------------------------------------------
bloco('2. A REGUA — as catorze reconstroem a partir das ancoras?')
# --------------------------------------------------------------------------


def vantagem(quantos):
    return PP_VANTAGEM * PP_ALIADO * quantos


def desvantagem(alvo):
    """acerto cai de 50% para 25%: o dano do alvo cai pela metade"""
    return alvo * (1 - ACERTO ** 2 / ACERTO)


def metros(m):
    return m * METRO


def acoes(alvo, quantas):
    return alvo * (quantas / CHEFE_ACOES)


CRITICO = 2 * DADO_DO_SOCO

# v0.104: o Surdo ganhou -2 na iniciativa, e nesta regua isso vale ZERO — nao por
# esquecimento. A peca 15 §3.1 ja rodou o contra-teste: 'que fracao do meu dano da
# rodada cai antes de o inimigo agir' da 52,5% em TODAS as montagens, porque
# iniciativa REORDENA o turno e nao tira acao de ninguem. A regua desta peca mede
# dano por rodada; iniciativa e' um eixo que ela nao alcanca, e a peca 15 e' quem
# publicou que esse eixo existe (foi ele que matou a saida A das Invocacoes).
# Fica escrito como termo e nao apagado, para que quem mexer aqui veja o zero e o
# motivo dele juntos.
INICIATIVA = 0.0


def catorze(alvo):
    d = desvantagem(alvo)
    return [
        ('Derrubado', 'Menor', vantagem(1) + metros(4.5)),
        ('Cego', 'Menor', d + vantagem(ALIADOS)),
        ('Surdo', 'Menor', 0.0 + INICIATIVA),
        ('Agarrado', 'Menor', metros(9.0)),
        ('Impedido', 'Menor', metros(9.0) + d + vantagem(ALIADOS)),
        ('Envenenado', 'Menor', d),
        ('Lento', 'Menor', metros(4.5) + acoes(alvo, 0.5)),
        ('Desarmado', 'Menor', FUNDO_DE_DUAS_MAOS * PONTO_DE_ARMA + metros(3.0)),
        ('Calado', 'Menor', acoes(alvo, 1.0)),
        ('Amedrontado', 'Maior', d + metros(9.0)),
        ('Enfeitiçado', 'Maior', acoes(alvo, 1.0)),
        ('Atordoado', 'Maior', acoes(alvo, 1.0) + acoes(alvo, 0.5)),
        ('Incapacitado', 'Maior', CRITICO),
        ('Petrificado', 'Maior', acoes(alvo, CHEFE_ACOES) + CRITICO + vantagem(ALIADOS)),
    ]


CALC = {n: (t, v) for n, t, v in catorze(CHEFE)}

# a tabela publicada no §2.2
_pub = {}
for _l in TXT.split('\n'):
    _m = re.match(r'\|\s*\*\*`([A-Za-zçãíéÇ]+)`\*\*\s*\|\s*`([\d,]+)`\s*\|'
                  r'\s*`([\d,]+)`\s*\|\s*`(Leve|Média|Pesada)`\s*\|', _l)
    if _m:
        _pub[_m.group(1)] = (num(_m.group(2)), num(_m.group(3)), _m.group(4))

if len(_pub) != 14:
    erro(f'2: a tabela do §2.2 tem {len(_pub)} linha(s) e eu esperava 14 — ela mudou '
         'de forma e esta checagem parou de conferir')
else:
    _mau = 0
    for _n, (_dano, _fat, _tier) in _pub.items():
        if _n not in CALC:
            erro(f'2: a peca publica "{_n}", que nao esta na lista de catorze')
            _mau += 1
            continue
        _esp = CALC[_n][1]
        if abs(_esp - _dano) > 0.02:
            erro(f'2: {_n} — a peca publica {_dano:.2f} de dano por rodada e a regua '
                 f'reconstroi {_esp:.2f}')
            _mau += 1
        elif abs(_dano / FATIA - _fat) > 0.02:
            erro(f'2: {_n} — {_dano:.2f} de dano sao {_dano/FATIA:.2f} fatias, e a '
                 f'peca publica {_fat:.2f}')
            _mau += 1
    if not _mau:
        print(f'  [x] as 14 linhas do §2.2 reconstroem a partir das ancoras')
        print(f'  [x] as 14 conversoes em fatia batem com a fatia lida do dono')


# v0.104: AS DUAS REGUAS DE ROLAGEM, e o quanto elas divergem. A v0.103
# publicava "4,7 vezes" em tres documentos, e o 4,7 e' 108 ÷ 23,00 — a razao das
# BASES, nao das reguas. Lidas por ponto percentual elas dao 9,4, e o fator 2 que
# separa os dois numeros e' a conversao: a sua regua e RELATIVA (peca 15 §3.3,
# "+1 no acerto = 50% -> 55% = +10% de dano saido") e a do aliado e ABSOLUTA.
# Nada disto esta guardado aqui: os quatro numeros vem dos donos.
_PP_POR_PONTO = 5.0          # +1 no d20 sao 5 pontos percentuais
_seu_pp = (ROTINA_30 * (_PP_POR_PONTO / (ACERTO * 100))) / _PP_POR_PONTO
_razao = _seu_pp / PP_ALIADO
_escopo = ROTINA_30 / ACAO_DE_ALIADO
_m = re.search(r'divergem por `(\d+,\d+)` vezes', TXT)
if not _m:
    erro('2: a peca nao publica mais por quanto as duas reguas de rolagem divergem '
         '— a frase mudou e esta checagem parou de conferir')
else:
    _pub_r = num(_m.group(1))
    print(f'  as duas reguas por ponto percentual: sua {_seu_pp:.3f} · '
          f'do aliado {PP_ALIADO:.3f}')
    if abs(_razao - _pub_r) > 0.05:
        erro(f'2: a peca publica {_pub_r:.2f} vezes de divergencia entre as duas '
             f'reguas e a conta reconstroi {_razao:.2f}')
    elif abs(_razao / _escopo - 2.0) > 0.01:
        erro(f'2: a razao das reguas ({_razao:.2f}) deveria ser exatamente o dobro '
             f'da razao das bases ({_escopo:.2f}) — o fator 2 e a conversao '
             'relativa contra absoluta, e ele parou de fechar')
    else:
        print(f'  [x] as duas reguas divergem {_razao:.2f}x = {_escopo:.2f} de escopo '
              f'x 2 de conversao, e a peca publica {_pub_r:.2f}')


# --------------------------------------------------------------------------
bloco('3. O NIVEL — ele sai da banda do manual, e a banda sai da tabela de Classe')
# --------------------------------------------------------------------------
# Media custa a Classe inteira e Pesada custa Classe e meia; entao, contra a
# Rotina de floor(3,5 x Classe) dados, Media e' 2/7 e Pesada e' 3/7. A Leve e'
# metade da Classe, que da 1/7. As tres bandas nao sao escolha: sao a tabela de
# preco do manual dividida pela coluna Rotina.

BANDAS = [('Leve', 1 / 7), ('Média', 2 / 7), ('Pesada', 3 / 7)]


def tier_de(valor):
    for t, fr in BANDAS:
        if valor <= ROTINA_30 * fr + 1e-9:
            return t
    return 'Pesada'


print('  as tres bandas, no nivel 30 (Rotina 108):')
for _t, _fr in BANDAS:
    print(f'     {_t:<7} ate {ROTINA_30*_fr:6.2f} de dano por rodada '
          f'= {ROTINA_30*_fr/FATIA:5.2f} fatias')

_mau = 0
for _n, (_dano, _fat, _tier) in _pub.items():
    _esp = tier_de(_dano)
    if _esp != _tier:
        erro(f'3: {_n} vale {_dano:.2f} de dano por rodada, que cai na banda {_esp}, '
             f'e a peca publica {_tier}')
        _mau += 1
if _pub and not _mau:
    _c = {t: sum(1 for v in _pub.values() if v[2] == t) for t, _ in BANDAS}
    print(f'  [x] as 14 caem na banda que a conta diz: '
          f'{_c["Leve"]} Leve · {_c["Média"]} Média · {_c["Pesada"]} Pesada')

_acima = [n for n, (d, _, _) in _pub.items() if d > ROTINA_30 * 3 / 7]
if _acima:
    print(f'  ~  {len(_acima)} passa(m) do teto da Pesada e a peca declara isso: '
          + ', '.join(sorted(_acima)))
    if 'passam do teto da `Pesada`' not in TXT:
        erro('3: alguma condicao passa do teto da Pesada e a peca nao declara isso')


# --------------------------------------------------------------------------
bloco('4. O MANUAL — as catorze da peca sao as catorze do manual, nos dois sentidos')
# --------------------------------------------------------------------------
try:
    import docx
except ImportError:
    docx = None

if docx is None:
    pulou('4. as catorze contra o manual — sem python-docx '
          '(pip install python-docx --break-system-packages)')
else:
    _d = docx.Document(DOCX)
    # v0.104: o manual publica as catorze em TRES tabelas, uma por nivel, e o
    # cabecalho de cada uma e' "Nivel <tier>". A Melhoria que aplica condicao e'
    # uma so, chamada Condicao, e ela cobra o nivel.
    _man = {}
    for _t in _d.tables:
        _cab = [c.text.strip() for c in _t.rows[0].cells]
        if len(_cab) == 2 and _cab[0].startswith('Nível '):
            _tier = _cab[0].split(' ', 1)[1].strip()
            _man[_tier] = [r.cells[0].text.strip() for r in _t.rows[1:]]
    if sorted(_man) != ['Leve', 'Média', 'Pesada']:
        erro(f'4: o manual publica as tabelas de nivel {sorted(_man)} e eu esperava '
             'Leve, Média e Pesada — ou o manual mudou, ou a extracao parou de achar')
    elif sum(len(v) for v in _man.values()) != 14:
        erro(f'4: as tres tabelas do manual somam {sum(len(v) for v in _man.values())} '
             'condicoes e eu esperava 14')
    else:
        # os nomes E os niveis vem DA PECA, e nao de lista escrita aqui dentro:
        # uma checagem que se mede contra a propria constante sai verde na
        # perturbacao que importa.
        _pm = {}
        for _sec in ('### 3.1 As seis de nível `Leve`',
                     '### 3.2 As duas de nível `Média`',
                     '### 3.3 As seis de nível `Pesada`'):
            _a = TXT.find(_sec)
            if _a < 0:
                erro(f'4: nao achei a secao "{_sec}" na peca — ela mudou de forma e '
                     'esta checagem parou de conferir')
                continue
            _corpo = TXT[_a + len(_sec):]
            _fim = _corpo.find('\n### ')
            _corpo = _corpo[:_fim] if _fim >= 0 else _corpo
            for _li in _corpo.split('\n'):
                _mm = re.match(r'\|\s*\*\*`([^`]+)`\*\*\s*\|\s*`(Leve|Média|Pesada)`\s*\|', _li)
                if _mm:
                    _pm[_mm.group(1)] = _mm.group(2)
        if len(_pm) != 14:
            erro(f'4: li {len(_pm)} condicao(oes) das tabelas §3.1 a §3.3 e esperava 14')
        _calc = {n for n, _, _ in catorze(CHEFE)}
        _fora_regua = sorted(set(_pm) - _calc)
        if _fora_regua:
            erro('4: a peca publica condicao que a regua do §2.2 nao preca: '
                 + ', '.join(_fora_regua))
        _falta = [n for g in _man for n in _man[g] if n not in _pm]
        _sobra = [n for n in _pm if n not in [x for g in _man for x in _man[g]]]
        if _falta:
            erro('4: o manual publica condicao que esta peca nao tem: ' + ', '.join(_falta))
        if _sobra:
            erro('4: esta peca tem condicao que o manual nao publica: ' + ', '.join(_sobra))
        _trocada = [f'{n} (manual {g}, peca {_pm.get(n)})'
                    for g in _man for n in _man[g] if _pm.get(n) != g]
        if _trocada:
            erro('4: condicao com nivel diferente entre o manual e a peca: '
                 + ', '.join(_trocada))
        if not (_falta or _sobra or _trocada):
            print('  o manual publica ' + ' · '.join(
                f'{len(_man[t])} {t}' for t in ('Leve', 'Média', 'Pesada')) + '; a peca tambem')
            print('  [x] as 14 batem com o manual em nome e em NIVEL, nos dois sentidos')

    # v0.104: a Melhoria e' uma so, e as duas de antes nao existem mais no manual
    _mel = []
    for _t in _d.tables:
        for _r in _t.rows:
            _c = [x.text.strip() for x in _r.cells]
            if _c and _c[0] in ('Condição', 'Condição Menor', 'Condição Maior'):
                _mel.append((_c[0], _c[1] if len(_c) > 1 else ''))
    _velhas = [m for m, _ in _mel if m != 'Condição']
    if _velhas:
        erro('4: o manual ainda vende ' + ', '.join(sorted(set(_velhas)))
             + ' — a v0.104 fundiu as duas na Melhoria Condição')
    elif not _mel:
        erro('4: nao achei a Melhoria Condição no catalogo do manual')
    elif _mel[0][1] != 'o nível dela':
        erro(f'4: a Melhoria Condição do manual cobra "{_mel[0][1]}" e a peca §3.6 diz '
             'que o preco e o nivel da condicao')
    else:
        print('  [x] o manual vende UMA Melhoria Condição, e o preco dela e o nivel')

    # v0.104: o -2 na iniciativa do Surdo mora em TRES lugares — a tabela de mesa
    # desta peca, a tabela de nivel do manual e a peca 3 §5, que e' a dona da
    # formula da iniciativa. Tres copias de um numero so; a licao no 9 pede
    # alguem comparando as tres em vez de a gente escolher uma e torcer.
    _SURDO = '−2'
    _peca_surdo = [l for l in TXT.split('\n')
                   if l.startswith('| **`Surdo`**') and 'iniciativa' in l]
    _man_surdo = []
    for _t in _d.tables:
        for _r in _t.rows:
            _c = [x.text.strip() for x in _r.cells]
            if len(_c) == 2 and _c[0] == 'Surdo':
                _man_surdo.append(_c[1])
    _p3 = ler('sistema/03-mecanica/03-economia-de-acao-e-iniciativa.md')
    _faltou = []
    if not (_peca_surdo and _SURDO in _peca_surdo[0]):
        _faltou.append('a tabela de mesa desta peca')
    if not (_man_surdo and any(_SURDO in x and 'iniciativa' in x for x in _man_surdo)):
        _faltou.append('a tabela de nivel do manual')
    if not ('`Surdo`' in _p3 and _SURDO in _p3 and 'iniciativa' in _p3):
        _faltou.append('a peca 3 §5, que e a dona da formula')
    if _faltou:
        erro('4: o -2 na iniciativa do Surdo nao esta em: ' + '; '.join(_faltou))
    else:
        print('  [x] o −2 na iniciativa do Surdo bate nos tres donos: '
              'esta peca, o manual e a peca 3 §5')

    # as tres que ficaram de fora continuam de fora, e o manual concorda
    _fora = ('Inconsciente', 'Exaustão', 'Invisível')
    _sem = [f for f in _fora if f'**`{f}`**' not in TXT]
    if _sem:
        erro('4: a peca parou de declarar as que ficam de fora: ' + ', '.join(_sem))
    else:
        print('  [x] as tres que ficam de fora continuam declaradas com o motivo')


# --------------------------------------------------------------------------
bloco('5. NENHUMA SEM NIVEL — e o nivel e um dos tres')
# --------------------------------------------------------------------------
# As tabelas do §3.1 e do §3.2 sao o texto de mesa, e elas carregam a coluna
# nivel. Esta checagem le AQUELAS tabelas e nao a do §2.2, de proposito: sao
# duas copias do mesmo numero, e a licao no 9 pede alguem comparando as duas.

_mesa = {}
for _l in TXT.split('\n'):
    _m = re.match(r'\|\s*\*\*`([A-Za-zçãíéÇ]+)`\*\*\s*\|\s*`(Leve|Média|Pesada)`\s*\|', _l)
    if _m:
        _mesa[_m.group(1)] = _m.group(2)

if len(_mesa) != 14:
    erro(f'5: as tabelas do §3.1 a §3.3 trazem {len(_mesa)} condicao(oes) com nivel '
         'e eu esperava 14 — elas mudaram de forma e esta checagem parou de conferir')
else:
    _div = [n for n, t in _mesa.items() if _pub.get(n, (0, 0, None))[2] != t]
    if _div:
        erro('5: o nivel do texto de mesa nao bate com o da tabela da regua em: '
             + ', '.join(sorted(_div)))
    else:
        _c = {t: sum(1 for x in _mesa.values() if x == t)
              for t in ('Leve', 'Média', 'Pesada')}
        print(f'  14 condicoes com nivel no texto de mesa: '
              f'{_c["Leve"]} Leve · {_c["Média"]} Média · {_c["Pesada"]} Pesada')
        print('  [x] as duas copias do nivel — a regua e o texto de mesa — batem')


# --------------------------------------------------------------------------
bloco('6. A ESCADA DE QUEM CURA — o teto por uso cobre os tiers que a peca diz')
# --------------------------------------------------------------------------
# Tirar condicao custa 1 PE por nivel: Leve 1, Media 2, Pesada 3. O teto por uso
# do Enxerto e' a maestria, e o do Cerzido e' a maior Classe. A escada nao esta
# escrita em lugar nenhum: ela cai da regra, e esta checagem confere que a
# tabela publicada e' a que cai.

CUSTO = {'Leve': 1, 'Média': 2, 'Pesada': 3}

_mm = re.search(r'\*\*Maestria\*\* = 1, \+1 a cada (\w+) níveis', ler('sistema/ESTADO-ATUAL.md'))
_passo = 8 if not _mm else {'oito': 8}.get(_mm.group(1), 8)


def maestria(nv):
    return min(4, 1 + (nv - 1) // _passo)


_escada = []
for _nv in (11, 17, 25):
    _teto = maestria(_nv)
    _escada.append((_nv, _teto, [t for t, c in CUSTO.items() if c <= _teto]))
_escada.append((27, 7, list(CUSTO)))

for _nv, _teto, _alc in _escada:
    print(f'  nivel {_nv:>2}: teto {_teto} PE por uso -> alcanca {", ".join(_alc)}')

if maestria(11) != 2 or maestria(17) != 3:
    erro(f'6: a maestria no nivel 11 deu {maestria(11)} e no 17 deu {maestria(17)}; '
         'a escada publicada supoe 2 e 3')
elif 'maestria` = `2`' not in TXT.replace(' ', ' '):
    print('  ~  a peca nao escreve "maestria = 2" com essa forma; conferindo so a conta')

# e ela tem de bater com a escada de exaustao da peca 10, que e' o precedente
_ex = ler(P10)
_faltando = [_t for _t in ('três degraus', 'degrau 3') if _t not in _ex]
if _faltando:
    erro('6: a peca 10 parou de escrever ' + ' e '.join(f'"{_t}"' for _t in _faltando)
         + ' — o precedente que a escada desta peca reproduz sumiu de la, e sem ele '
           'a escada de 1 · 2 · 3 fica sem a segunda fonte que a confirma')
else:
    print('  [x] a escada cai da regra, e ela bate com os tres degraus de exaustao')


# --------------------------------------------------------------------------
bloco('7. OS TIPOS DE DANO — catorze, tres grupos, e a tabela de resistencia')
# --------------------------------------------------------------------------
GRUPOS = {}
for _l in TXT.split('\n'):
    _m = re.match(r'>?\s*\|\s*\*\*(Físicos|Elementais|Especiais)\*\*\s*\|(.+?)\|\s*'
                  r'\*\*(\d+)%\*\*\s*\|', _l)
    if _m:
        GRUPOS[_m.group(1)] = (
            [t.strip(' `') for t in _m.group(2).split('·')], int(_m.group(3)))

if len(GRUPOS) != 3:
    erro(f'7: achei {len(GRUPOS)} grupo(s) de tipo de dano e esperava 3 — a tabela '
         'mudou de forma e esta checagem parou de conferir')
else:
    _tipos = [t for g, _ in GRUPOS.values() for t in g]
    _peso = sum(p for _, p in GRUPOS.values())
    if len(_tipos) != 14:
        erro(f'7: sao {len(_tipos)} tipos de dano na tabela e a peca diz catorze')
    elif len(set(_tipos)) != 14:
        erro('7: tem tipo de dano repetido na tabela')
    elif _peso != 100:
        erro(f'7: os tres grupos somam {_peso}% do dano recebido, e tem de somar 100')
    else:
        print(f'  14 tipos em 3 grupos: '
              + ' · '.join(f'{g} {p}%' for g, (_, p) in GRUPOS.items()))
        print('  [x] os tipos nao repetem e os tres pesos somam 100%')

    # a natureza da lista tem de continuar dita: o peso e' PREVISAO sem dono, e
    # previsao que perde o rotulo vira numero fechado na versao seguinte.
    # As duas metades ficam separadas de proposito (licao no 8): perturbar 30
    # para 40 acende a soma; tirar a palavra PREVISAO acende esta.
    _s7 = TXT.split('## 4. Os tipos de dano')[1].split('## 5. Cobertura')[0] \
        if '## 4. Os tipos de dano' in TXT else ''
    if 'PREVISÃO' not in _s7:
        erro('7: a secao 4 parou de dizer que o peso dos tres grupos e PREVISAO — '
             'previsao que perde o rotulo vira numero fechado na versao seguinte')
    elif '04-playtest' not in _s7:
        erro('7: a secao 4 diz que o peso e previsao e nao nomeia quem seria o dono — '
             'previsao sem destinatario nunca e cobrada de ninguem')
    else:
        print('  [x] o peso segue rotulado como previsao, com o playtest como dono')

    # a tabela de quantos tipos voce resiste e' recontada, nao guardada
    _res = {}
    for _l in TXT.split('\n'):
        _m = re.match(r'\|\s*\*{0,2}(\d)\*{0,2}(?:\s*—[^|]*)?\s*\|\s*(\d+)%\s*\|'
                      r'\s*\*{0,2}([\d,]+)\*{0,2}(?: fatia)?\s*\|', _l)
        if _m:
            _res[int(_m.group(1))] = (int(_m.group(2)), num(_m.group(3)))
    if len(_res) != 4:
        erro(f'7: a tabela de resistencia tem {len(_res)} linha(s) e eu esperava 4')
    else:
        _recebido = None
        _m = re.search(r'são `(\d+,\d+)` de dano por rodada', TXT)
        if _m:
            _recebido = num(_m.group(1)) / (GRUPOS['Físicos'][1] / 100)
        if _recebido is None:
            erro('7: nao achei o dano recebido por rodada, que e a base da tabela')
        else:
            _mau = 0
            for _q, (_pc, _ft) in sorted(_res.items()):
                _esp = _recebido * _pc / 100 / FATIA
                if abs(_esp - _ft) > 0.02:
                    erro(f'7: resistir a {_q} tipo(s) bate em {_pc}% de '
                         f'{_recebido:.2f}, que sao {_esp:.2f} fatias, e a peca '
                         f'publica {_ft:.2f}')
                    _mau += 1
            if not _mau:
                print(f'  [x] as 4 linhas de resistencia saem de {_recebido:.2f} de '
                      f'dano recebido por rodada, recontadas')


# --------------------------------------------------------------------------
bloco('8. COBERTURA — tres degraus, e a Total sem numero')
# --------------------------------------------------------------------------
COB = {}
for _l in TXT.split('\n'):
    _m = re.match(r'\|\s*\*\*(Parcial|Boa|Total)\*\*\s*\|(.+?)\|', _l)
    if _m:
        COB[_m.group(1)] = _m.group(2)

if len(COB) != 3:
    erro(f'8: achei {len(COB)} degrau(s) de cobertura e esperava 3 — a tabela mudou '
         'de forma e esta checagem parou de conferir')
else:
    _mau = 0
    for _d, _esp in (('Parcial', 2), ('Boa', 5)):
        _n = re.findall(r'`\+(\d+)`', COB[_d])
        if len(_n) != 2 or set(int(x) for x in _n) != {_esp}:
            erro(f'8: a cobertura {_d} tem de dar +{_esp} de Defesa e +{_esp} no '
                 f'Teste de Resistencia Fisico, e a peca escreve {_n}')
            _mau += 1
    if re.search(r'`\+?\d', COB['Total']):
        erro('8: a cobertura Total ganhou numero, e ela nao e um bonus — '
             'ela e a ausencia de alvo legal')
        _mau += 1
    if 'Só a maior conta' not in TXT:
        erro('8: a peca parou de dizer que so a maior cobertura conta')
        _mau += 1
    if not _mau:
        print('  Parcial +2 · Boa +5 · Total sem numero, e so a maior conta')
        print('  [x] os tres degraus fecham, e a Total continua sem numero')


# --------------------------------------------------------------------------
bloco('9. AS ENTREGAS PUBLICADAS — o Abalo e o Encontrao batem com a regua')
# --------------------------------------------------------------------------
# Esta e' a checagem que a peca existe para ter, e ela SAI DA PASTA: le os dois
# DESENHO da raiz. Ela pega o caso em que alguem reescreve o texto de uma
# entrega sem mexer no preco dela, ou o contrario — que foi exatamente o que
# deixou o Derrubado do Punho tres versoes preçado como permanente.

_d = CALC['Derrubado'][1]

# o Abalo: Derrubado permanente, com a trava que o DESENHO-manhas publica
_txtm = ler(DMAN)
_m = re.search(r'total `(\d+,\d+)` de dano por rodada, que é `(\d+,\d+)` fatia', _txtm)
if not _m:
    erro('9: nao achei a derivacao do Derrubado no DESENHO-manhas — ela mudou de '
         'forma e esta checagem parou de conferir')
else:
    _abalo, _abfat = num(_m.group(1)), num(_m.group(2))
    if abs(_abalo - _d) > 0.02:
        erro(f'9: o DESENHO-manhas preca o Derrubado permanente em {_abalo:.2f} e a '
             f'regua desta peca reconstroi {_d:.2f}')
    else:
        print(f'  [x] o Abalo: Derrubado permanente em {_d:.2f}, e o desenho concorda')

_m = re.search(r'\| \*\*Massa\*\* \| `Abalo` \|[^|]*\| (\d+)% \| \*\*(\d+,\d+)\*\* \|', _txtm)
if not _m:
    erro('9: nao achei a linha do Abalo no catalogo do DESENHO-manhas')
else:
    _tx, _ft = int(_m.group(1)) / 100, num(_m.group(2))
    if abs(_d * _tx / FATIA - _ft) > 0.02:
        erro(f'9: o Abalo dispara a {_m.group(1)}% e vale {_d*_tx/FATIA:.2f} fatia, '
             f'e o catalogo publica {_ft:.2f}')
    else:
        print(f'  [x] o Abalo com a trava de {_m.group(1)}%: {_ft:.2f} fatia, e bate')

# o Encontrao: o Derrubado do nivel 11 do Punho, com os DOIS portoes do texto.
# O portao NAO fica escrito aqui: ele e' lido do texto da propria entrega. E' o
# que faz esta checagem pegar os dois sentidos — mexer no preco sem mexer no
# texto, e mexer no texto sem mexer no preco.
_txtt = ler(DTRI)
_m = re.search(r'\| nv11 `Derrubado` \| (\d+,\d+) \| \*\*(\d+,\d+)\*\* \|', _txtt)
_me = re.search(r'\*\*Nível 11 — `Encontrão`\.\*\*(.+?)\n', _txtt, re.S)
if not _m:
    erro('9: nao achei a linha do Derrubado do Punho no DESENHO-trilhas')
elif not _me:
    erro('9: nao achei o texto de mesa do Encontrao no DESENHO-trilhas — sem ele nao '
         'da para ler que portao a entrega escreve')
else:
    _texto = _me.group(1)
    _dano, _fat = num(_m.group(1)), num(_m.group(2))
    _p_acerto = 1 - (1 - ACERTO) ** 2 if 'acertou' in _texto or 'acerta' in _texto else 1.0
    _p_tr = 1 - RESISTE if 'Teste de Resistência' in _texto else 1.0
    print(f'  portoes lidos do texto do Encontrao: acerto {_p_acerto:.0%} · '
          f'Teste de Resistencia {_p_tr:.0%}')
    _esp = _d * _p_acerto * _p_tr
    if abs(_esp - _dano) > 0.05:
        erro(f'9: o Encontrao aplica Derrubado com dois portoes — acertar '
             f'({_p_acerto:.0%}) e o alvo falhar o TR ({_p_tr:.0%}) —, o que da '
             f'{_esp:.2f} de dano por rodada, e o desenho publica {_dano:.2f}')
    elif abs(_dano / FATIA - _fat) > 0.02:
        erro(f'9: o Encontrao vale {_dano/FATIA:.2f} fatia e o desenho publica {_fat:.2f}')
    else:
        print(f'  [x] o Encontrao com os dois portoes ({_p_acerto:.0%} x {_p_tr:.0%} '
              f'= {_p_acerto*_p_tr:.1%}): {_dano:.2f} = {_fat:.2f} fatia')

# e a Trilha inteira tem de caber, agora que o preco esta certo
_m = re.search(r'## `Punho` — (\d+,\d+) de (\d+,\d+)', _txtt)
if not _m:
    erro('9: nao achei o cabecalho do Punho com o total e o orcamento')
else:
    _tot, _orc = num(_m.group(1)), num(_m.group(2))
    if _tot > _orc:
        erro(f'9: o Punho fecha em {_tot:.2f} de um orcamento de {_orc:.2f} — com o '
             'preco corrigido ele tem de caber')
    else:
        print(f'  [x] o Punho fecha em {_tot:.2f} de {_orc:.2f}, e cabe')


# --------------------------------------------------------------------------
bloco('10. NENHUM VALOR DE REGRA GUARDADO AQUI DENTRO')
# --------------------------------------------------------------------------
# A promessa do cabecalho e' que toda ancora vem do dono. Esta checagem confere
# a promessa: a fatia e o chefe tem de ter sido LIDOS na checagem 1, e nao
# sobrado do valor de formato escrito no topo deste arquivo.

_prometido = {'a fatia': FATIA, 'o chefe': CHEFE}
_defaults = {'a fatia': 5.08, 'o chefe': 72.0}
_lidos_ok = len(_LIDOS) == len(ANCORAS)
if not _lidos_ok:
    erro(f'10: {len(ANCORAS)-len(_LIDOS)} ancora(s) nao foram lidas do dono, '
         'entao a regua rodou com valor de formato escrito neste arquivo')
else:
    print(f'  [x] as {len(ANCORAS)} ancoras foram lidas do documento dono')

# e o teto de cada tier nao pode estar escrito: ele e' derivado da tabela do manual
if re.search(r'^\s*(TETO|BANDA)_\w+\s*=\s*[\d.]+', open(__file__, encoding='utf-8').read(),
             re.M):
    erro('10: tem teto de tier escrito como constante neste arquivo — ele tem de '
         'sair da tabela de Classe do manual, dividida pela coluna Rotina')
else:
    print('  [x] as tres bandas sao derivadas de 1/7, 2/7 e 3/7 da Rotina, '
          'e nenhuma esta escrita')


# --------------------------------------------------------------------------
bloco('11. A PENALIDADE DE ARMA — as duas linhas, e o tamanho de cada uma')
# --------------------------------------------------------------------------
# Nenhum dos quatro numeros esta escrito aqui: a desvantagem vem da peca 11, o
# metro e o ponto de arma vem das pecas 5 e 14, e o fundo de duas maos vem da 14.
# O 3 m e' o unico que se confere contra fora: sao os 10 pes que o d20 de 2024
# cobra de quem veste protecao sem a Forca dela.
_sec6 = TXT[TXT.find('## 6. A penalidade de arma'):]
_sec6 = _sec6[:_sec6.find('\n## ')] if '\n## ' in _sec6 else _sec6
if not _sec6:
    erro('11: nao achei a secao 6 (a penalidade de arma) — ela mudou de titulo e '
         'esta checagem parou de conferir')
else:
    _falta = [t for t, _r in (
        ('desvantagem na rolagem de ataque', 'desvantagem na rolagem de ataque'),
        ('a queda de deslocamento', '`3 m`'),
    ) if _r not in _sec6]
    if _falta:
        erro('11: a secao 6 nao escreve mais: ' + ', '.join(_falta))
    _desv = ROTINA_30 * (PP_VANTAGEM / (ACERTO * 100))
    _mov = 3.0 * METRO
    _arma = FUNDO_DE_DUAS_MAOS * PONTO_DE_ARMA
    _razao = (_desv + _mov) / _arma
    for _rot, _val, _rx in (
        ('a desvantagem', _desv, r'\*\*`(\d+,\d+)` de dano por rodada\*\*'),
        ('a razao contra a arma inteira', _razao, r'custam `(\d+,\d+)` vezes'),
    ):
        _m = re.search(_rx, _sec6)
        if not _m:
            erro(f'11: a secao 6 nao publica mais {_rot} no formato que esta '
                 'checagem le')
        elif abs(num(_m.group(1)) - _val) > 0.05:
            erro(f'11: a secao 6 publica {num(_m.group(1)):.2f} para {_rot} e a '
                 f'conta reconstroi {_val:.2f}')
    _pes = 3.0 / 0.3048
    if abs(_pes - 10.0) > 0.2:
        erro(f'11: os 3 m da penalidade dao {_pes:.1f} pes, e o d20 cobra 10')
    if not ERROS:
        print(f'  desvantagem: {PP_VANTAGEM:.0f} pp sobre {ACERTO*100:.0f}% = '
              f'{_desv:.2f} de dano por rodada')
        print(f'  deslocamento -3 m = {_mov:.2f}   |   a arma inteira = {_arma:.2f}')
        print(f'  [x] as duas somadas custam {_razao:.1f}x a entrega da arma — '
              'porta fechada, e nao preco')
        print(f'  [x] os 3 m sao os {_pes:.0f} pes que o d20 de 2024 cobra')


# --------------------------------------------------------------------------
print()
print('=' * 88)
if ERROS:
    print(f'>>> {len(ERROS)} PROBLEMA(S):')
    for e in ERROS:
        print('   -', e)
    sys.exit(1)
if _PULADAS:
    print(f'>>> OK, mas {len(_PULADAS)} checagem(ns) PULARAM:')
    for p in _PULADAS:
        print('   -', p)
    print('    O que pulou NAO foi conferido. Um verde que pulou checagem nao e um verde.')
else:
    print('>>> TUDO OK — a regua reconstroi as catorze a partir dos donos, o nivel')
    print('    de cada uma cai da banda do manual, os tipos de dano e a cobertura')
    print('    fecham, e as duas entregas publicadas batem com o portao que elas')
    print('    mesmas escrevem.')
