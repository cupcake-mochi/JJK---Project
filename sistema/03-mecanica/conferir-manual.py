#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Confere o MANUAL contra o PROJETO. E a direcao que faltava.

O conferir-nomes.py olha projeto -> manual: "esse nome que eu batizei ja significa
alguma coisa la?". Ninguem olhava o contrario: "o manual usa alguma palavra que
este sistema nao tem?".

E por isso que o "Bonus de Treinamento" da Passiva Reforco sobreviveu ate a v0.25
e o "Habilidade / Sabedoria" da Restricao Fraqueza sobreviveu ate a v0.26. Os dois
sao vocabulario de outro sistema vivo dentro do manual, e nenhum dos cinco
validadores olhava para la: o pac7.py e o v7.py conferem NUMERO, e o conferir-nomes
confere a outra direcao.

Quatro checagens:
  1. VOCABULARIO ORFAO — palavra de outro sistema que este aqui nao tem. Sabedoria
     e Carisma fundiram em Essencia; o nosso bonus de treino se chama Maestria;
     "Habilidade" aqui e atributo; "Grau" e patente e nao tamanho de feitico.
     Tolerancia zero: se aparecer, alguem escreveu texto novo sem olhar o lado de ca.
  2. OS CINCO ATRIBUTOS — nenhum outro nome de atributo pode aparecer no manual.
  3. TERMO MECANICO SEM DEFINICAO — palavra que o manual usa como se fosse termo
     definido e nunca explica. As que ja existem estao declaradas aqui com motivo,
     no mesmo padrao do conferir-pericias; uma nova FALHA.
  4. OS NUMEROS COMPARTILHADOS, COM DONO DECLARADO — a tabela de PE, a de inimigo
     e a curva de Rotina aparecem nos dois lados. Se divergirem, o projeto mente em
     silencio.
     ATENCAO AO QUE ESTA CHECAGEM *NAO* DIZ. Ela nao diz que o manual esta certo.
     Os limitadores e exemplos dele foram calibrados quando o sistema em volta era
     outro; eles servem de continuidade, nao de lei. Divergencia aqui e' um pedido
     de DECISAO — qual dos dois lados muda —, e nao um veredito de que o projeto
     errou. Por isso cada numero carrega um dono declarado logo abaixo.

Roda sem argumento. Sai com codigo 1 se algo quebrar.
Sem python-docx, as quatro checagens sao PULADAS com aviso, em vez de falhar.
"""

import os
import re
import sys

FALHAS = []
AVISOS = []


def erro(msg):
    FALHAS.append(msg)
    print(f'  !! {msg}')


def aviso(msg):
    AVISOS.append(msg)
    print(f'  ~~ {msg}')


def bloco(t):
    print()
    print('=' * 88)
    print(t)
    print('=' * 88)


# --------------------------------------------------------------------------
AQUI = os.path.dirname(os.path.abspath(__file__))
RAIZ = os.path.abspath(os.path.join(AQUI, '..', '..'))
DOCX = os.path.join(RAIZ, 'manual', 'Fundamento-MANUAL-v7.docx')

try:
    import docx  # noqa: F401
except ImportError:
    print('~~ python-docx nao instalado. As quatro checagens foram PULADAS.')
    print('   pip install python-docx --break-system-packages')
    sys.exit(0)

if not os.path.exists(DOCX):
    print(f'~~ manual nao encontrado em {DOCX}. As quatro checagens foram PULADAS.')
    sys.exit(0)

import docx as _docx
_D = _docx.Document(DOCX)

LINHAS = []          # (origem, texto)
for i, p in enumerate(_D.paragraphs):
    if p.text.strip():
        LINHAS.append((f'paragrafo {i}', p.text))
for ti, t in enumerate(_D.tables):
    for ri, r in enumerate(t.rows):
        txt = ' | '.join(c.text for c in r.cells)
        if txt.strip():
            LINHAS.append((f'tabela {ti}, linha {ri}', txt))

TUDO = '\n'.join(t for _, t in LINHAS)

print(f'Manual lido: {len(_D.paragraphs)} paragrafos, {len(_D.tables)} tabelas, '
      f'{len(LINHAS)} linhas com texto.')


# --------------------------------------------------------------------------
bloco('1. VOCABULARIO ORFAO — palavra de outro sistema viva no manual')

# termo -> (o que este sistema usa no lugar, por que a troca aconteceu)
ORFAOS = {
    r'\bSabedoria\b':            ('Essencia', 'Sabedoria e Carisma fundiram em Essencia'),
    r'\bCarisma\b':              ('Essencia', 'Sabedoria e Carisma fundiram em Essencia'),
    r'\bHabilidade\b':           ('atributo', 'aqui os cinco se chamam atributos'),
    r'B[oô]nus de Treinamento':  ('Maestria', 'o nosso numero de treino e a Maestria'),
    r'Classe de Armadura':       ('Defesa', 'Defesa = 10 + Destreza + protecao'),
    r'Percep[çc][aã]o Passiva':  ('a pericia Percepcao', 'nao existe valor passivo aqui'),
    r'Profici[êe]ncia':          ('Maestria', 'o treino e binario e o numero e a Maestria'),
    r'\bGrau\b':                 ('Classe', 'Grau e a patente do feiticeiro desde a v0.20'),
    r'\bEscala\b':               ('Classe', 'nome de rascunho do tamanho do feitico'),
    r'\bPot[êe]ncia\b':          ('Classe', 'nome de rascunho do tamanho do feitico'),
    r'\btruque\b':               ('feitico de Classe 0', 'vocabulario de outro sistema'),
    r'\bcantrip\b':              ('feitico de Classe 0', 'vocabulario de outro sistema'),
    r'\bn[ií]vel de conjurador\b': ('nivel', 'aqui existe um nivel so'),
}

achou_orfao = False
for rx, (subst, motivo) in ORFAOS.items():
    hits = [(o, t) for o, t in LINHAS if re.search(rx, t)]
    if hits:
        achou_orfao = True
        nome = rx.replace(r'\b', '').replace('[oô]', 'o').replace('[êe]', 'e') \
                 .replace('[çc]', 'c').replace('[aã]', 'a').replace('[ií]', 'i')
        erro(f'"{nome}" aparece {len(hits)}x no manual. Aqui isso e "{subst}" — {motivo}')
        for o, t in hits[:3]:
            print(f'        {o}: {t[:150]}')
if not achou_orfao:
    print(f'  Nenhum dos {len(ORFAOS)} termos de outro sistema aparece no manual.')
    print('  Esta e a checagem que teria pego o Bonus de Treinamento na v0.24 e o')
    print('  Habilidade/Sabedoria da Fraqueza tres versoes antes de a v0.26 achar.')


# --------------------------------------------------------------------------
bloco('2. TESTE NOMEADO PELO ATRIBUTO, EM VEZ DO TESTE DE RESISTENCIA')

TRS = ['Fisico', 'Vigor', 'Intelecto', 'Espirito']
ATRIBUTOS = ['Forca', 'Destreza', 'Constituicao', 'Inteligencia', 'Essencia']
print('  Aqui existem quatro Testes de Resistencia — ' + ' · '.join(TRS) + ' — e cada um')
print('  tem um atributo por baixo. O manual nao pode chamar um teste pelo ATRIBUTO:')
print('  "teste de Constituicao" nao diz se e o TR Vigor ou outra coisa, e o TR Fisico')
print('  usa Forca OU Destreza, entao a traducao nem e um para um.\n')

# "teste de X" / "Teste de Resistencia de X", onde X e' nome de atributo
RX_TESTE = (r'[Tt]este(?:s)? (?:de Resist[êe]ncia )?de (?:uma )?'
            r'(For[çc]a|Destreza|Constitui[çc][aã]o|Intelig[êe]ncia|Ess[êe]ncia|'
            r'Sabedoria|Carisma|Habilidade)')
achou_teste = False
for o, t in LINHAS:
    for m in re.finditer(RX_TESTE, t):
        achou_teste = True
        erro(f'{o}: teste nomeado pelo atributo — "{m.group(0)}". '
             f'Aqui os testes se chamam {", ".join(TRS)}')
        print(f'        {t[:170]}')
if not achou_teste:
    print('  Nenhum teste do manual e nomeado pelo atributo.')

# uma regra de preco que liste atributos so pode listar os cinco daqui
DE_FORA = ['Sabedoria', 'Carisma', 'Astucia', 'Aparencia', 'Percepcao Passiva']
fora = [x for x in DE_FORA
        if re.search(r'\b' + x.replace('c', '[çc]').replace('e', '[êe]') + r'\b', TUDO)]
if fora:
    erro(f'atributo que nao existe neste sistema citado no manual: {", ".join(fora)}')
else:
    print('\n  E nenhum atributo de fora aparece. Uma regra de preco que cite atributo')
    print('  (como a Restricao Fraqueza) so pode citar os cinco daqui.')

print('\n  Contagem, so para leitura — a palavra tambem aparece como Tema e como prosa:')
for nome, rx in zip(ATRIBUTOS, [r'For[çc]a', r'Destreza', r'Constitui[çc][aã]o',
                                r'Intelig[êe]ncia', r'Ess[êe]ncia']):
    print(f'    {nome:<16}{len(re.findall(rx, TUDO)):>3}x')
print('    (Vontade e Forca aparecem tambem como Tema, nos grupos "Mente e alma" e')
print('     "Forca e movimento". Tema nao e atributo, e por isso a checagem olha para')
print('     a frase "teste de X" em vez de para a palavra solta.)')


# --------------------------------------------------------------------------
bloco('3. TERMO MECANICO USADO E NUNCA DEFINIDO')

# termo -> (regex, motivo pelo qual ele esta declarado aqui em vez de falhar)
INDEFINIDOS_ACEITOS = {
    'cobertura leve': (
        r'cobertura leve',
        'a Passiva Afinidade fura cobertura leve. Cobertura e conceito de tabuleiro '
        'e mora na peca de equipamento, que ainda nao existe. Reavaliar quando ela sair.'),
    'inimigo fraco': (
        r'[Ii]nimigos? fracos?',
        'a Passiva Peso da Presenca so pega inimigo fraco. Depende do bestiario, '
        'que sai da matematica de inimigo do proprio manual.'),
}

# Termos que o manual IMPORTA do projeto, de proposito. Eles sao a direcao
# contraria do problema que este validador existe para pegar: em vez de vocabulario
# de outro sistema vazando para dentro, e vocabulario DESTE sistema entrando porque
# uma peca precisou dele. Cada um so entra aqui com o lugar onde o manual o define.
#
# O teste generico de "esta definido" (a constante DEFINE) aceita qualquer frase
# que contenha a palavra "e" — e quase toda frase em portugues contem. Isso basta
# para os termos da checagem acima, que sao raros e aparecem poucas vezes. NAO
# basta aqui: um termo importado aparece em varias linhas, e uma delas vai casar
# por acidente. Entao cada importado declara o PROPRIO padrao de definicao.
IMPORTADOS_DO_PROJETO = {
    'refino': (
        r'\brefino\b',
        r'[Oo] \*?\*?refino\*?\*? é',
        'a Expansao de Dominio (v7.7) tem gate de refino, desconto de refino e '
        'duracao por refino. O manual define o termo na caixa "REFINO, EM UMA LINHA" '
        'da secao 7 e nao usa ele em mais lugar nenhum.'),
}

# termos que EXIGEM definicao no manual: se aparecerem sem uma linha que os
# explique, e sem estarem declarados acima, falha
EXIGEM_DEFINICAO = {
    'dano fisico':  r'dano f[ií]sico',
    'resistencia':  r'resist[êe]ncia (ao|a) (seu )?tipo de dano|sem resist[êe]ncia',
}

DEFINE = r'\b[ée]\b|significa|quer dizer|considera-se|chamamos|:\s*metade|metade do dano'

for nome, rx in EXIGEM_DEFINICAO.items():
    usos = [(o, t) for o, t in LINHAS if re.search(rx, t)]
    if not usos:
        print(f'  {nome:<16} 0 usos — nao esta no manual')
        continue
    defs = [t for _, t in usos if re.search(DEFINE, t)]
    if defs:
        print(f'  {nome:<16} {len(usos)} uso(s), definido')
    elif nome in INDEFINIDOS_ACEITOS:
        print(f'  {nome:<16} {len(usos)} uso(s), indefinido — ACEITO')
    else:
        erro(f'"{nome}" e usado {len(usos)}x no manual e nunca definido. '
             f'Dois mestres leem diferente, e o preco da peca que o usa depende disso')
        for o, t in usos[:3]:
            print(f'        {o}: {t[:150]}')

print()
print('  Termos IMPORTADOS do projeto — o manual usa, e tem que definir:')
for nome, (rx, rx_def, motivo) in IMPORTADOS_DO_PROJETO.items():
    usos = [(o, t) for o, t in LINHAS if re.search(rx, t)]
    if not usos:
        aviso(f'"{nome}" esta declarado como termo importado e nao aparece mais no '
              'manual — a declaracao virou peso morto')
        continue
    defs = [t for _, t in usos if re.search(rx_def, t)]
    if not defs:
        erro(f'"{nome}" e termo do PROJETO, aparece {len(usos)}x no manual e nunca e '
             'definido la. Quem le so o manual nao sabe o que ele e, e o gate que '
             'depende dele vira numero sem unidade')
    else:
        print(f'    {nome} ({len(usos)}x, definido no manual)')
        print(f'      motivo: {motivo}')

print()
print('  Indefinidos ACEITOS, com motivo declarado:')
for nome, (rx, motivo) in INDEFINIDOS_ACEITOS.items():
    n = len([1 for _, t in LINHAS if re.search(rx, t)])
    if n == 0:
        aviso(f'"{nome}" esta declarado como indefinido aceito e nao aparece mais '
              f'no manual — a declaracao virou peso morto')
    else:
        print(f'    {nome} ({n}x)')
        print(f'      motivo: {motivo}')


# --------------------------------------------------------------------------
bloco('4. OS NUMEROS COMPARTILHADOS, E QUEM MANDA EM CADA UM')
print('  Estes numeros aparecem nos DOIS lados. Se divergirem, o projeto mente em')
print('  silencio — mas divergir nao quer dizer que o projeto errou. Os limitadores')
print('  do manual foram calibrados quando o sistema em volta era outro: eles servem')
print('  de continuidade, e nao de lei.\n')
DONO = {
    'PE': ('o PROJETO', 'nada exige que o Emanador tenha 6 de PE por nivel. O que exige e '
           'que a coluna "quantas vezes voce lanca" diga a verdade sobre a ficha. Mudou o '
           '6? Regere a coluna. O numero e nosso; a coluna e a consequencia'),
    'inimigo': ('o PLAYTEST', 'esta e a unica das tres que afirma alguma coisa sobre o '
                'mundo: que um combate dura ~3,5 rodadas. A trava de vida inteira da peca 1 '
                'foi calibrada contra ela. Ninguem e dono ate alguem jogar'),
    'Classe 0': ('o MANUAL', 'ele tem tabela propria — 2d8 . 3d8 . 4d8 . 5d8 . 6d8 por '
                 'faixa de nivel — e ate a v0.79 nenhum documento do projeto e nenhum '
                 'validador abriam ela. A peca 6 precava ele em 4,50 fixo, que nao existe '
                 'no manual. Ele e a QUARTA tabela compartilhada, e era a unica sem dono'),
    'Rotina': ('o MANUAL', 'ela nao e uma medida, e a DEFINICAO de "quanto dano por rodada '
               'e normal". Nao ha verdade fora dela — o projeto compara tudo contra ela, '
               'inclusive ela mesma. Mudar a Rotina e mudar a regua, e reprecifica tudo'),
}
for k, (quem, motivo) in DONO.items():
    print(f'  {k:<10} dono: {quem}')
    print(f'             {motivo}')
print()

# (rotulo, o que o projeto assume, funcao que extrai do .docx)
def _tabela_com(cabecalho_contem):
    for t in _D.tables:
        hdr = ' | '.join(c.text.strip() for c in t.rows[0].cells)
        if all(k in hdr for k in cabecalho_contem):
            return t
    return None


ok_num = True

# 4a. PE por nivel do conjurador — a peca 1 secao 5.3 diz que a formula veio daqui
t = _tabela_com(['PE total'])
if t is None:
    erro('nao achei a tabela de PE total no manual — a peca 1, secao 5.3 diz que a '
         'formula do PE maximo vem dela')
    ok_num = False
else:
    print(f"  {'nivel':<8}{'PE do manual':<16}{'6 x nivel':<12}bate?")
    for r in t.rows[1:]:
        cel = [c.text.strip() for c in r.cells]
        try:
            nv, pe = int(cel[0]), int(cel[1])
        except ValueError:
            continue
        bate = pe == 6 * nv
        print(f'  {nv:<8}{pe:<16}{6*nv:<12}{"sim" if bate else "NAO"}')
        if not bate:
            ok_num = False
            erro(f'PE: o manual diz {pe} no nivel {nv} e a formula do projeto da {6*nv}. '
                 f'DONO: {DONO["PE"][0]} — entao o normal e regerar a coluna do manual, '
                 f'e nao mudar a peca 1. Se a decisao for a outra, mude os dois')

# 4b. dano de chefe e capanga — o conferir-atributos.py tem essa tabela dentro
CHEFE_NO_PROJETO = {5: 15, 10: 26, 15: 38, 20: 49, 25: 61, 30: 72}
t = _tabela_com(['Chefe', 'Capanga'])
if t is None:
    erro('nao achei a tabela de inimigos no manual — o conferir-atributos.py copia '
         'o dano de chefe dela para medir rodadas sob foco')
    ok_num = False
else:
    print(f"\n  {'nivel':<8}{'chefe no manual':<18}{'no projeto':<14}bate?")
    for r in t.rows[1:]:
        cel = [c.text.strip() for c in r.cells]
        try:
            nv, dano = int(cel[0]), int(cel[3])
        except (ValueError, IndexError):
            continue
        esperado = CHEFE_NO_PROJETO.get(nv)
        bate = esperado == dano
        print(f'  {nv:<8}{dano:<18}{str(esperado):<14}{"sim" if bate else "NAO"}')
        if not bate:
            ok_num = False
            erro(f'inimigo: o manual diz {dano} de dano de chefe no nivel {nv} e o '
                 f'conferir-atributos.py assume {esperado}. DONO: {DONO["inimigo"][0]} — '
                 f'esta e a tabela que promete ~3,5 rodadas, e mexer nela move a trava '
                 f'de vida inteira da peca 1')

# 4c. a coluna Rotina — a peca 6 usa ela para aprovar ataque extra e invocacao
ROTINA_NO_PROJETO = {1: 13, 2: 31, 3: 45, 4: 63, 5: 76, 6: 94, 7: 108}
t = _tabela_com(['Rotina'])
if t is None:
    erro('nao achei a coluna Rotina no manual — a peca 6 pendura nela o ataque '
         'extra e o orcamento de invocacao')
    ok_num = False
else:
    print(f"\n  {'Classe':<8}{'Rotina no manual':<20}{'no projeto':<14}bate?")
    for r in t.rows[1:]:
        cel = [c.text.strip() for c in r.cells]
        try:
            cl = int(cel[1])
        except (ValueError, IndexError):
            continue
        m = re.search(r'=\s*(\d+)', cel[2])
        if not m:
            continue
        dano = int(m.group(1))
        esperado = ROTINA_NO_PROJETO.get(cl)
        bate = esperado == dano
        print(f'  {cl:<8}{dano:<20}{str(esperado):<14}{"sim" if bate else "NAO"}')
        if not bate:
            ok_num = False
            erro(f'Rotina da Classe {cl}: o manual diz {dano} e a peca 6 assume '
                 f'{esperado}. DONO: {DONO["Rotina"][0]} — ela e a regua, nao uma medida. '
                 f'Mudar a Rotina reprecifica o golpe canalizado, o ataque extra e a '
                 f'invocacao de uma vez')

if ok_num:
    print('\n  Os tres conjuntos batem. Enquanto baterem, nenhuma das dez decisoes do')
    print('  projeto penduradas neles precisa ser reaberta.')


# --------------------------------------------------------------------------
# 4d. A ROTINA POR NIVEL DA PECA 6 SS3 CONTRA A COLUNA DO MANUAL
#
# Escrita na v0.60, e ela existe por um vao que a 4c acima NAO cobre.
#
# A 4c confere a coluna Rotina do .docx contra um dicionario escrito aqui. Ela sai
# VERDE com a peca 6 publicando qualquer coisa, porque ela nunca abre a peca 6. Foi
# por esse vao que o 81 e o 126 sobreviveram catorze versoes: os dois moram na MESMA
# tabela do manual, em OUTRAS colunas — 'Feitico num alvo' da Classe 6 e 'Somando
# alvos' da Classe 7. Numero que veio da coluna errada da tabela certa passa por
# qualquer varredura que so procure se o numero existe no manual.
#
# NADA DE VALOR FICA ESCRITO AQUI, e isso vale para o mapa tambem: a faixa de nivel
# de cada Classe sai da coluna 'Nivel' da PROPRIA tabela do manual. Se o manual
# reagrupar as faixas, esta checagem acompanha sozinha.
#
# O QUE TEM DE ACENDER: trocar qualquer Rotina de qualquer tabela da peca 6 por
# outro numero da mesma linha do manual. Contra-teste: trocar pela Rotina certa de
# OUTRO nivel tambem tem de acender, senao a checagem so confere "existe no manual".
print()
print('  4d. a Rotina por NIVEL da peca 6 SS3 contra a coluna do manual')

_p6 = os.path.join(AQUI, '06-caminhos-e-trilhas.md')
_t = _tabela_com(['Rotina'])
if not os.path.exists(_p6):
    erro('nao achei a peca 6 para conferir a Rotina por nivel')
elif _t is None:
    erro('nao achei a coluna Rotina no manual para conferir a peca 6 SS3')
else:
    _i_rot = [n for n, c in enumerate(_t.rows[0].cells)
              if c.text.strip() == 'Rotina']
    _faixas = []
    for _r in _t.rows[1:]:
        _cel = [c.text.strip() for c in _r.cells]
        _n = re.findall(r'\d+', _cel[0])
        _m = re.search(r'=\s*(\d+)', _cel[_i_rot[0]]) if _i_rot else None
        if len(_n) >= 2 and _m:
            _faixas.append((int(_n[0]), int(_n[1]), int(_m.group(1))))

    def _rotina_do_nivel(nv):
        for _a, _b, _v in _faixas:
            if _a <= nv <= _b:
                return _v
        return None

    _txt = open(_p6, encoding='utf-8').read()
    # varre TABELA por TABELA, e so as que declaram Rotina no cabecalho
    _achados, _cab, _col = [], None, None
    for _lin in _txt.splitlines():
        _s = _lin.strip()
        if not _s.startswith('|'):
            _cab, _col = None, None
            continue
        _cel = [x.strip() for x in _s.strip('|').split('|')]
        if _cab is None:
            _r = [n for n, c in enumerate(_cel) if 'Rotina' in c]
            if _r:
                _cab, _col = _cel, _r[0]
            continue
        if not _cel or not _cel[0].isdigit():
            continue
        _m = re.match(r'\**\s*(\d+)', _cel[_col]) if _col < len(_cel) else None
        if _m:
            _achados.append((int(_cel[0]), int(_m.group(1))))

    if not _achados:
        erro('a peca 6 nao publica Rotina por nivel em tabela nenhuma — ou o '
             'formato mudou, e esta checagem parou de conferir em silencio')
    else:
        print(f"    {'nivel':<8}{'peca 6 diz':<13}{'manual, pela faixa':<21}bate?")
        _vistos = set()
        for _nv, _val in _achados:
            _esp = _rotina_do_nivel(_nv)
            _bate = _esp == _val
            if (_nv, _val) not in _vistos:
                _vistos.add((_nv, _val))
                print(f'    {_nv:<8}{_val:<13}{str(_esp):<21}'
                      f'{"sim" if _bate else "NAO"}')
            if not _bate:
                _onde = []
                for _r in _t.rows[1:]:
                    _c = [x.text.strip() for x in _r.cells]
                    for _j, _x in enumerate(_c):
                        if _j >= 2 and re.search(rf'=\s*{_val}$', _x):
                            _onde.append(f'Classe {_c[1]} · coluna "'
                                         f'{_t.rows[0].cells[_j].text.strip()}"')
                erro(f'peca 6: no nivel {_nv} ela publica Rotina {_val} e o manual '
                     f'diz {_esp}. DONO: {DONO["Rotina"][0]}. '
                     + (f'O {_val} existe no manual, mas em ' + ' e '.join(_onde)
                        + ('. Coluna certa, LINHA errada.'
                           if all('"Rotina"' in _o for _o in _onde)
                           else ' — coluna errada da tabela certa.')
                        if _onde else
                        f'O {_val} nao existe em coluna nenhuma daquela tabela.'))
        if all(_rotina_do_nivel(n) == v for n, v in _achados):
            print(f'    As {len(_vistos)} linhas de Rotina da peca 6 saem da coluna '
                  f'certa do manual,')
            print('    e a faixa de cada Classe foi lida do .docx em vez de escrita aqui.')


# --------------------------------------------------------------------------
# 4e. A ROTINA RECONSTROI DO PROPRIO MANUAL — e ela NUNCA foi "feitico + Classe 0"
#
# Escrita na v0.80. Ela existe porque a peca 6 SS3 passou de v0.14 ate aqui
# explicando a coluna Rotina com uma frase que nao reconstroi de nada:
# "a coluna Rotina do Fundamento ja e feitico + Classe 0". Ela nao e.
#
# A Rotina e o MEIO EXATO entre as duas colunas vizinhas da mesma tabela:
#     Feitico num alvo = 3 x Classe dados   (regra de ouro no 2: "para nos pontos")
#     Somando alvos    = 4 x Classe dados   (o teto)
#     Rotina           = floor(3,5 x Classe) dados
# Bate nas SETE Classes, com zero parametro livre.
#
# NADA DE VALOR ESCRITO AQUI: as tres contagens de dados de cada linha saem do .docx.
#
# O QUE TEM DE ACENDER: mexer na contagem de dados de qualquer Rotina do manual.
# CONTRA-TESTE: a leitura velha — "num alvo + Classe 0" — tem de dar DIFERENTE da
# Rotina em pelo menos uma Classe. Se ela desse igual, esta checagem estaria
# aprovando as duas leituras ao mesmo tempo e nao provaria nada.
print()
print('  4e. a Rotina reconstroi como o meio entre "num alvo" e "somando alvos"')

def _dados(txt):
    """soma todos os NdX de uma celula: '21d8 + 3d8 = 108' -> 24"""
    return sum(int(n) for n in re.findall(r'(\d+)d\d+', txt))

_tr = _tabela_com(['Rotina', 'Feitico num alvo']) or _tabela_com(['Rotina', 'Feitiço num alvo'])
_t0 = None
for _t in _D.tables:
    _l = [[c.text.strip() for c in r.cells] for r in _t.rows]
    if _l and _l[0][0].startswith('Seu n') and any(r[0] == 'Dano' for r in _l):
        _t0 = _l
        break

if _tr is None:
    erro('nao achei a tabela da curva no manual — a Rotina e a regua de tudo')
elif _t0 is None:
    erro('nao achei a tabela de dano do Classe 0 no manual. DONO: o MANUAL. '
         'Sem ela nao da para conferir a leitura velha da Rotina')
else:
    _cab = [c.text.strip() for c in _tr.rows[0].cells]
    _ir = next(n for n, c in enumerate(_cab) if c == 'Rotina')
    _ia = next(n for n, c in enumerate(_cab) if 'num alvo' in c)
    _is = next(n for n, c in enumerate(_cab) if 'Somando' in c)

    # o Classe 0 por FAIXA DE NIVEL, lido do .docx
    _nv0 = [int(x) for x in _t0[0][1:]]
    _dd0 = [int(re.match(r'(\d+)', x).group(1)) for x in _t0[2][1:]]

    def _c0_dados(nv):
        _v = _dd0[0]
        for _n, _x in zip(_nv0, _dd0):
            if nv >= _n:
                _v = _x
        return _v

    print(f"    {'Classe':<8}{'num alvo':<11}{'somando':<10}{'o meio':<9}"
          f"{'Rotina':<9}{'bate?':<7}{'num alvo + C0'}")
    _velha_bate_sempre = True
    for _r in _tr.rows[1:]:
        _cel = [c.text.strip() for c in _r.cells]
        if not _cel[1].isdigit():
            continue
        _cl = int(_cel[1])
        _da, _ds, _drot = _dados(_cel[_ia]), _dados(_cel[_is]), _dados(_cel[_ir])
        _meio = (_da + _ds) // 2
        _niv_da_classe = int(re.findall(r'\d+', _cel[0])[0])
        _velha = _da + _c0_dados(_niv_da_classe)
        if _velha != _drot:
            _velha_bate_sempre = False
        _bate = (_meio == _drot) and (_drot == 3 * _cl + _cl // 2)
        print(f'    {_cl:<8}{str(_da)+"d8":<11}{str(_ds)+"d8":<10}'
              f'{str(_meio)+"d8":<9}{str(_drot)+"d8":<9}'
              f'{"sim" if _bate else "NAO":<7}{str(_velha)+"d8"}')
        if not _bate:
            erro(f'Rotina da Classe {_cl}: o manual publica {_drot}d8, e o meio entre '
                 f'"num alvo" ({_da}d8) e "somando alvos" ({_ds}d8) da {_meio}d8 '
                 f'(= floor(3,5 x Classe) = {3*_cl + _cl//2}d8). DONO: '
                 f'{DONO["Rotina"][0]}. A Rotina e a regua: mudar ela reprecifica o '
                 f'golpe canalizado, o ataque extra, a invocacao e as quinze Trilhas')

    if _velha_bate_sempre:
        erro('CONTRA-TESTE FALHOU: "num alvo + Classe 0" deu igual a Rotina em TODAS '
             'as Classes. Entao esta checagem nao separa a leitura certa da leitura '
             'que viveu de v0.14 a v0.79, e ela nao esta provando nada')
    else:
        print('    Contra-teste: "num alvo + Classe 0" NAO reproduz a Rotina. A frase')
        print('    "a Rotina ja e feitico + Classe 0" morreu na v0.80, e esta linha e')
        print('    o que impede ela de voltar.')


# --------------------------------------------------------------------------
# 4f. O CLASSE 0 TEM DONO, E A PECA 6 PAROU DE INVENTAR UM
#
# Escrita na v0.80. O manual publica o dano de um Classe 0 numa tabela propria —
# 2d8 . 3d8 . 4d8 . 5d8 . 6d8 por faixa de nivel — e ate a v0.79 NENHUM documento
# do projeto e NENHUM validador abriam ela. A peca 6 SS3 precava o Classe 0 em 4,50
# em todo nivel, que e um numero que nao aparece em lugar nenhum do manual.
#
# O estrago: a coluna "conjurador" da peca 6 saia 5 pontos alta em todo nivel, e o
# vao "fisico - conjurador" — que paga o degrau do nivel 7 dos cinco Caminhos, o
# nivel 2 do Arremate e o empate em +6% — saia 4/5/6/7 quando ele e 9/10/11/12.
#
# A REGRA APLICADA, e ela e separada do limite de design de proposito:
#     a coluna "conjurador" da peca 6 e o feitico SOZINHO ("Feitico num alvo"),
#     porque um Classe 0 gasta a Acao Padrao e nao cabe junto do feitico grande.
# O LIMITE DE DESIGN e outro: o vao tem de ser positivo e crescer com o nivel,
#     porque ele e um golpe simples.
#
# NADA DE VALOR ESCRITO AQUI: o feitico por Classe e a faixa de cada Classe saem
# do .docx; os niveis publicados saem da peca 6.
#
# O QUE TEM DE ACENDER: somar qualquer coisa de volta na coluna conjurador.
print()
print('  4f. a coluna "conjurador" da peca 6 contra o feitico sozinho do manual')

if not os.path.exists(_p6):
    erro('nao achei a peca 6 para conferir a linha de base do SS3')
elif _tr is None:
    erro('nao achei a tabela da curva no manual para conferir a peca 6')
else:
    _faixa_cl = []
    _feitico = {}
    for _r in _tr.rows[1:]:
        _cel = [c.text.strip() for c in _r.cells]
        if not _cel[1].isdigit():
            continue
        _n = re.findall(r'\d+', _cel[0])
        _cl = int(_cel[1])
        _feitico[_cl] = int(re.search(r'=\s*(\d+)', _cel[_ia]).group(1))
        if len(_n) >= 2:
            _faixa_cl.append((int(_n[0]), int(_n[1]), _cl))

    def _classe_do_nivel(nv):
        for _a, _b, _c in _faixa_cl:
            if _a <= nv <= _b:
                return _c
        return None

    # a tabela de base do SS3: nivel | Rotina | conjurador | fisico
    _linhas_base, _dentro = [], False
    for _lin in open(_p6, encoding='utf-8').read().splitlines():
        _s = _lin.strip()
        if not _s.startswith('|'):
            _dentro = False
            continue
        _cel = [x.strip() for x in _s.strip('|').split('|')]
        if not _dentro:
            if len(_cel) == 4 and 'Rotina' in _cel[1] and 'conjurador' in _cel[2]:
                _dentro = True
            continue
        if len(_cel) == 4 and _cel[0].isdigit():
            try:
                _linhas_base.append(tuple(int(re.match(r'\**\s*(-?\d+)', x).group(1))
                                          for x in _cel))
            except AttributeError:
                pass

    if not _linhas_base:
        erro('a peca 6 nao publica mais a tabela "nivel | Rotina | conjurador | '
             'fisico" — ou o formato mudou, e esta checagem parou em silencio')
    else:
        print(f"    {'nv':<5}{'conjurador':<12}{'feitico sozinho':<17}"
              f"{'fisico':<9}{'vao':<7}bate?")
        _vao_ant = None
        for _nv, _rot, _cj, _fi in _linhas_base:
            _cl = _classe_do_nivel(_nv)
            _esp = _feitico.get(_cl)
            _vao = _fi - _cj
            _bate = _esp == _cj
            print(f'    {_nv:<5}{_cj:<12}{str(_esp):<17}{_fi:<9}{_vao:<7}'
                  f'{"sim" if _bate else "NAO"}')
            if not _bate:
                erro(f'peca 6 SS3: no nivel {_nv} ela publica conjurador {_cj} e o '
                     f'feitico sozinho da Classe {_cl} e {_esp}. DONO do Classe 0: '
                     f'o MANUAL, na tabela de dano do Classe 0. A diferenca de '
                     f'{_cj - _esp} e o Classe 0 fantasma de 4,50 que viveu de v0.14 '
                     f'a v0.79 — um Classe 0 gasta a Acao Padrao e nao cabe junto do '
                     f'feitico grande')
            if _vao <= 0:
                erro(f'peca 6 SS3: o vao no nivel {_nv} deu {_vao}. Ele e um golpe '
                     f'simples, entao ele e positivo — e ele paga o degrau do nivel 7 '
                     f'dos cinco Caminhos')
            if _vao_ant is not None and _vao < _vao_ant:
                erro(f'peca 6 SS3: o vao encolheu do nivel anterior para o {_nv} '
                     f'({_vao_ant} -> {_vao}). Ele e um golpe simples e o golpe '
                     f'simples so cresce')
            _vao_ant = _vao

    # A frase morta, guardada POR LINHA e nao pelo arquivo inteiro.
    # Linha de citacao (">") e nota em italico ("*texto*") sao historia e podem
    # conter a frase — o projeto guarda o erro em vez de apagar.
    #
    # v0.81: o teste de historia estava ERRADO e deixava passar linha viva.
    # Ele aceitava qualquer linha comecando com "*", e "**negrito**" comeca com "*"
    # — e negrito no comeco da linha e o estilo dominante da prosa deste projeto.
    # Toda afirmacao viva em negrito era lida como nota historica.
    # Agora: ">" e historia, "*" sozinho e historia, "**" e AFIRMACAO VIVA.
    def _e_historica(_l):
        _s = _l.lstrip()
        if _s.startswith('>'):
            return True
        return _s.startswith('*') and not _s.startswith('**')

    _vivas, _historicas = [], 0
    for _n, _lin in enumerate(open(_p6, encoding='utf-8').read().splitlines(), 1):
        if 'feitiço + Classe 0' not in _lin:
            continue
        if _e_historica(_lin):
            _historicas += 1
        else:
            _vivas.append(_n)
    for _n in _vivas:
        erro(f'peca 6, linha {_n}: ela afirma que a Rotina e "feitico + Classe 0". '
             f'A checagem 4e prova que nao e — a Rotina e o meio entre bater num '
             f'alvo e espalhar, e o Classe 0 tem tabela propria no manual. Se for '
             f'nota historica, ela vai num bloco de citacao')
    if not _vivas:
        print(f'    A frase morta nao aparece viva em nenhuma linha '
              f'({_historicas} em nota historica, que e onde ela deve ficar).')

    # 4g. O NUMERO morto, e nao so a frase.
    # A frase morta era "a Rotina ja e feitico + Classe 0". O numero que ela
    # produzia era 4,50 de dano por Classe 0, e ele sobreviveu a v0.80 na SS5,
    # onde argumentava o PE do Bastiao — sem a frase, so o numero.
    # Guarda: nenhuma linha VIVA pode preçar um Classe 0 em 4,5 ou 4,50.
    # O dono do dano de um Classe 0 e o manual, e a tabela dele esta acima.
    _num, _num_hist = [], 0
    for _n, _lin in enumerate(open(_p6, encoding='utf-8').read().splitlines(), 1):
        if 'Classe 0' not in _lin:
            continue
        if not re.search(r'4,50?(?![0-9])', _lin):
            continue
        if _e_historica(_lin):
            _num_hist += 1
        else:
            _num.append(_n)
    print()
    print('  4g. o NUMERO morto do Classe 0 (4,50), e nao so a frase morta')
    for _n in _num:
        erro(f'peca 6, linha {_n}: ela preca um Classe 0 em 4,50. Esse numero nao '
             f'existe no manual — ele e o dano de UM d8, que e a regua de montar '
             f'feitico. O dano de um Classe 0 tem tabela propria no manual, e a '
             f'checagem 4f a le. Se for nota historica, ela vai num bloco de citacao')
    if not _num:
        print(f'    Nenhuma linha viva preca Classe 0 em 4,50 '
              f'({_num_hist} em nota historica).')


# --------------------------------------------------------------------------
# 4h. A FORMA DO ATAQUE EXTRA — e nao so o numero dele
#
# Escrita na v0.82 para guardar a forma "golpe SOLTO"; INVERTIDA na v0.147, quando
# o Mizuki reverteu aquela decisao. O motivo da inversao esta na peca 6 SS3.1 e ele
# e concreto: com o golpe solto, o `Bote` — nivel 19 da `Estocada` — comprava por
# 2,46 fatias uma coisa que ja acontecia sozinha. Entrega preçada valendo zero.
#
# O QUE ELA GUARDA AGORA, e sao TRES metades independentes de proposito:
#   (a) a peca 6 declara que o ataque extra EXIGE a Acao de Atacar;
#   (b) nenhuma linha VIVA da peca 6 afirma o contrario — a forma velha nao pode
#       voltar por descuido, e a tabela historica dela mora num bloco de citacao;
#   (c) a VALVULA continua escrita: "a nao ser que uma habilidade diga o contrario".
#       Sem ela o `Bote` morre de novo, e a inversao inteira perde o motivo.
#
# A (c) e a que importa mais, e ela e nova. As duas primeiras guardam a forma; a
# terceira guarda o que a forma existe para permitir.
#
# O PRECO DESTA FORMA ESTA MEDIDO E ACEITO: dois golpes rendem 23 no nivel 30
# contra 27 de um Classe 0 gratis, entao a Acao de Atacar fica dominada pelo botao
# que toda ficha tem. A v0.82 recusou a forma por isso; a v0.147 a escolhe sabendo.
print()
print('  4h. a FORMA do ataque extra na peca 6 — exige a Acao de Atacar, com valvula')

if not os.path.exists(_p6):
    erro('nao achei a peca 6 para conferir a forma do ataque extra')
else:
    _txt6 = open(_p6, encoding='utf-8').read()

    def _historica_h(_l):
        _s = _l.lstrip()
        if _s.startswith('>'):
            return True
        return _s.startswith('*') and not _s.startswith('**')

    # (a) a declaracao afirmativa existe — e ela e a LINHA DE REGRA, nao o titulo
    #     da secao. O arnes pegou a primeira versao desta checagem passando no
    #     proprio titulo "O ataque extra EXIGE a Acao de Atacar": apagar a regra
    #     saia VERDE. Mesmo defeito da checagem 2 do conferir-alma.py.
    _regra6 = [l for l in _txt6.splitlines()
               if re.search(r'ganha um golpe simples por rodada', l, re.I)]
    _regra6 = _regra6[0] if _regra6 else ''
    _decl = _regra6 and re.search(r'exige a A[cç][aã]o de Atacar', _regra6, re.I)
    if not _decl:
        erro('peca 6: a forma do ataque extra nao esta declarada. Ela precisa dizer '
             'que ele EXIGE a Acao de Atacar — foi a inversao da v0.147, e o que ela '
             'existe para consertar e o `Bote` da `Estocada` valer zero')
    else:
        print('    [x] a peca 6 declara que o ataque extra exige a Acao de Atacar')

    # (b) a forma velha nao pode voltar viva
    _contra = []
    for _n, _lin in enumerate(_txt6.splitlines(), 1):
        if not re.search(r'ataque extra|golpe simples', _lin, re.I):
            continue
        if not re.search(r'n[aã]o exige a A[cç][aã]o de Atacar|golpe\s+SOLTO|'
                         r'golpe simples solto', _lin, re.I):
            continue
        if _historica_h(_lin):
            continue
        _contra.append(_n)
    for _n in _contra:
        erro(f'peca 6, linha {_n}: a forma VELHA do ataque extra voltou — golpe solto, '
             f'sem exigir a Acao de Atacar. Ela foi invertida na v0.147, e com ela o '
             f'`Bote` da `Estocada` volta a valer zero com preco de 2,46 fatias. Se for '
             f'nota historica, ela vai num bloco de citacao')
    if not _contra:
        print('    [x] a forma velha so aparece em nota historica')

    # (c) A VALVULA — e ela e o motivo de a inversao caber.
    #     Tambem na LINHA DE REGRA: o paragrafo que explica a valvula carrega uma
    #     segunda copia dela, e o arnes mostrou que apagar a valvula da regra saia
    #     verde por causa dessa copia.
    if not re.search(r'a n[aã]o ser que uma habilidade diga o contr[aá]rio', _regra6, re.I):
        erro('peca 6: sumiu a valvula "a nao ser que uma habilidade diga o contrario". '
             'Sem ela nenhuma entrega de Trilha consegue comprar a excecao, e o `Bote` '
             'volta a ser letra morta — que e exatamente o defeito que a v0.147 saiu '
             'para consertar')
    else:
        print('    [x] a valvula continua escrita: habilidade pode dizer o contrario')

    # (c) o gate do golpe do Arremate e do Coro continua escrito — a outra metade
    #     da forma que a propria peca disse que nenhum validador guardava
    if not re.search(r'A[cç][aã]o B[oô]nus.{0,200}A[cç][aã]o Padr[aã]o', _txt6, re.S):
        erro('peca 6: sumiu o gate do golpe do Arremate e do Coro — ele e Acao '
             'Bonus e so existe se a Acao Padrao daquele turno foi gasta no que a '
             'Trilha e. Sem o gate, uma Padrao solta conjura, golpeia e ainda '
             'sobra a Bonus')
    else:
        print('    [x] o gate do golpe do Arremate e do Coro continua escrito')


# --------------------------------------------------------------------------
# 4i. A PAREDE DO MANUAL, que a peca 11 SS6.6 copiou
#
# Escrita na v0.91. A peca 11 preca a vida da `Barreira Simples` comparando com a
# Melhoria `Anteparo` do manual — "uma parede com 10 x Classe de vida" — e essa
# frase virou a SEGUNDA copia do numero. Sem esta checagem, o manual podia mudar
# a Melhoria e a peca continuaria preçando contra o valor velho.
#
# NADA ESCRITO AQUI: os dois lados sao lidos, um do .docx e o outro da peca 11.
print()
print('  4i. a parede da Melhoria `Anteparo` contra a copia da peca 11')

_ant_manual = None
for _t in _D.tables:
    for _r in _t.rows:
        _cel = [c.text.strip() for c in _r.cells]
        if _cel and _cel[0].strip() == 'Anteparo':
            _mm = re.search(r'(\d+)\s*[×x]\s*Classe', ' | '.join(_cel))
            if _mm:
                _ant_manual = int(_mm.group(1))
if _ant_manual is None:
    erro('nao achei a Melhoria `Anteparo` no .docx — a peca 11 preca a `Barreira '
         'Simples` contra ela e a comparacao ficou sem dono')
else:
    _p11m = os.path.join(AQUI, '11-aptidoes-e-refino.md')
    with open(_p11m, encoding='utf-8') as _f:
        _t11m = _f.read()
    _mm = re.search(r'`(\d+) × Classe` de vida', _t11m)
    _ant_peca = int(_mm.group(1)) if _mm else None
    if _ant_peca is None:
        erro('a peca 11 parou de citar a parede do `Anteparo` — a vida da `Barreira '
             'Simples` deixou de ter contra o que ser comparada')
    elif _ant_peca != _ant_manual:
        erro(f'o manual diz que o `Anteparo` da {_ant_manual} x Classe de vida e a peca '
             f'11 copiou {_ant_peca} x Classe — as duas copias divergiram, e a peca '
             f'preca a `Barreira Simples` contra a copia dela')
    else:
        print(f'    [x] as duas dizem {_ant_manual} x Classe. No Classe 7 sao '
              f'{_ant_manual*7} de vida.')


# --------------------------------------------------------------------------
# 4j. A ESCADA DE FREQUENCIA DO EFEITO PROPRIO, que a peca 11 SS6.7 copiou
#
# Escrita na v0.92. A peca 11 listou "falta a regua do Efeito Proprio" por sessenta
# versoes, e o manual publica ela numa tabela: "Em quantas cenas por arco isso vai
# importar? Uma cena: Leve. Metade: Media. Quase toda: Pesada. Na duvida, Pesada."
#
# A peca 6.7 usa essa escada para dizer em que degrau de Classe Passiva uma
# `Aptidao Propria` cai — entao ela virou a segunda copia das tres faixas.
print()
print('  4j. a escada de frequencia do `Efeito Proprio` contra a copia da peca 11')

_freq = None
for _t in _D.tables:
    for _r in _t.rows:
        _cel = ' | '.join(c.text.strip() for c in _r.cells)
        if 'Em quantas cenas por arco' in _cel:
            _freq = _cel
if _freq is None:
    erro('nao achei no .docx a pergunta do `Efeito Proprio` — a peca 11 SS6.7 preca a '
         '`Aptidao Propria` contra ela e a escada ficou sem dono')
else:
    _tres = [(k, v) for k, v in (('Uma cena', 'Leve'), ('Metade', 'Média'),
                                 ('Quase toda', 'Pesada'))
             if f'{k}: {v}' in _freq]
    if len(_tres) != 3:
        erro(f'a escada do `Efeito Proprio` no .docx nao tem as tres faixas na ordem '
             f'esperada — achei {len(_tres)}. A peca 11 SS6.7 mapeia as tres para os '
             f'tres degraus de Classe Passiva, e o mapa quebrou')
    elif 'Na dúvida, Pesada' not in _freq:
        erro('o manual parou de dizer "Na duvida, Pesada" na escada do `Efeito '
             'Proprio` — e o desempate e o que faz a peca 11 SS6.7 RECUSAR a proposta '
             'em vez de aceitar')
    else:
        _p11f = os.path.join(AQUI, '11-aptidoes-e-refino.md')
        with open(_p11f, encoding='utf-8') as _f:
            _t11f = _f.read()
        _copia = re.findall(r'\| \*\*(uma|metade|quase toda)\*\* \| (Leve|Média|Pesada) \|',
                            _t11f)
        _esp = [('uma', 'Leve'), ('metade', 'Média'), ('quase toda', 'Pesada')]
        if _copia != _esp:
            erro(f'a peca 11 SS6.7 copiou a escada do `Efeito Proprio` como {_copia} e o '
                 f'manual diz {_esp} — as duas copias divergiram')
        else:
            print('    [x] as tres faixas batem: uma/Leve, metade/Media, quase toda/Pesada.')
            print('    [x] o desempate "Na duvida, Pesada" continua nos dois.')


# --------------------------------------------------------------------------
# 4k. A LISTA DE PASSIVAS POR CLASSE PASSIVA, que a peca 11 SS4 copiou
#
# Escrita na v0.107, e ela nasceu de uma divergencia achada a olho na revisao do
# livro: a peca 11 publicava `—` na linha da Classe Passiva 3 — NENHUMA Passiva —
# enquanto o manual lista tres ali (`Escama`, `Afinidade`, `Reserva Profunda`).
# A linha da 2 tambem estava curta: cinco de sete.
#
# A coluna existe como PROVA de que a escada de formato foi lida do manual e nao
# inventada na peca. Uma prova que diverge do que ela cita prova o contrario.
#
# `Regra Propria` e `Passiva Propria` sao `1 a 3` e ficam de fora dos dois lados:
# elas nao moram numa altura so.
#
# NADA ESCRITO AQUI: as tres linhas saem do .docx, e as tres saem da peca 11.
print()
print('  4k. as Passivas por Classe Passiva no .docx contra a copia da peca 11 SS4')

_pv_manual = {'1': [], '2': [], '3': []}
for _t in _D.tables:
    _cab = [c.text.strip() for c in _t.rows[0].cells] if _t.rows else []
    if _cab[:3] != ['Passiva', 'Classe', 'O que faz']:
        continue
    for _r in _t.rows[1:]:
        _cel = [c.text.strip() for c in _r.cells]
        if len(_cel) >= 2 and _cel[1] in _pv_manual:
            _pv_manual[_cel[1]].append(_cel[0])

if not all(_pv_manual.values()):
    erro('nao achei a lista de Passivas do .docx com as tres alturas preenchidas — '
         'achei ' + repr({k: len(v) for k, v in _pv_manual.items()}) + '. A peca 11 SS4 '
         'copia essa lista como prova da escada, e a comparacao ficou sem dono')
else:
    _p11p = os.path.join(AQUI, '11-aptidoes-e-refino.md')
    with open(_p11p, encoding='utf-8') as _f:
        _t11p = _f.read()
    _pv_peca = {}
    for _lin in _t11p.splitlines():
        _m = re.match(r'\|\s*\*\*([123])\*\*\s*\|[^|]*\|\s*([^|]+?)\s*\|\s*$', _lin)
        if _m:
            _pv_peca[_m.group(1)] = [x.strip(' `') for x in _m.group(2).split('·')]
    if set(_pv_peca) != {'1', '2', '3'}:
        erro('nao consegui ler as tres linhas da tabela de Classe Passiva da peca 11 SS4 '
             f'— achei {sorted(_pv_peca)}. Se a tabela mudou de forma, esta checagem '
             'parou de conferir')
    else:
        _ruim = False
        for _cl in ('1', '2', '3'):
            if _pv_peca[_cl] != _pv_manual[_cl]:
                _ruim = True
                erro(f'Classe Passiva {_cl}: o manual lista {_pv_manual[_cl]} e a peca 11 '
                     f'SS4 copiou {_pv_peca[_cl]} — as duas copias divergiram, e a coluna '
                     f'da peca existe justamente para provar que a escada saiu do manual')
        if not _ruim:
            print('    [x] ' + ' · '.join(f'Classe Passiva {c}: {len(_pv_manual[c])}'
                                          for c in ('1', '2', '3')) +
                  ' — as tres linhas batem, nome por nome e na mesma ordem.')

# guarda: `Regra Propria` e `Passiva Propria` sao `1 a 3` e nao podem aparecer em
# nenhuma das tres linhas dos dois lados. Se o manual passar a dar altura fixa a
# uma delas, a comparacao acima muda de forma e esta guarda acusa primeiro.
_flex = [_c[0] for _t in _D.tables for _c in
         ([[x.text.strip() for x in _r.cells] for _r in _t.rows])
         if len(_c) >= 2 and _c[1] == '1 a 3']
if sorted(_flex) != ['Passiva Própria', 'Regra Própria']:
    erro(f'as Passivas de altura flexivel do .docx mudaram: achei {sorted(_flex)} e '
         f'esperava a `Regra Própria` e a `Passiva Própria`. A checagem 4k deixa as '
         f'duas de fora das tres linhas, e essa exclusao deixou de valer')
else:
    print('    [x] a `Regra Própria` e a `Passiva Própria` continuam `1 a 3`, fora das tres.')


# --------------------------------------------------------------------------
bloco('5. O PORTAO DAS LINHAS DE CONTROLE — quem prende o alvo diz como solta')
# --------------------------------------------------------------------------
# v0.151. O `Cerca` passou nove versoes sendo a UNICA linha de Controle que
# prendia um alvo e nao dizia como aquilo acaba: o `Prende` cobra acao mais Teste
# de Resistencia, o `Anteparo` tem pontos de vida, o `Desarma o Feitico` tem
# portao de Classe, e as condicoes `Pesada` dao Teste de Resistencia no fim de
# cada turno. Ele nao tinha nada, e custava metade do `Prende`.
#
# ⚠ A classificacao NAO esta escrita aqui dentro. Quem escolhe as linhas que
# precisam de portao e' o texto do proprio manual: a linha fala do ALVO, e ela
# dura alem do instante. `Terreno` e `Anteparo` falam da area e da parede;
# `Puxa` e `Desarma o Feitico` acontecem e acabam. Nenhum dos quatro entra, e
# nenhum deles esta nomeado neste arquivo.

_tc = _tabela_com(['Melhoria', 'O que faz'])
_ctrl = []
if _tc is None:
    erro('5: nao achei a tabela de Melhorias no .docx — a checagem do portao de '
         'Controle ficou sem chao')
else:
    # o catalogo tem uma tabela por Familia, todas com o mesmo cabecalho: a de
    # Controle e' a que carrega a Melhoria `Condicao`.
    for _t5 in _D.tables:
        _h5 = ' | '.join(c.text.strip() for c in _t5.rows[0].cells)
        if 'Melhoria' not in _h5 or 'O que faz' not in _h5:
            continue
        _linhas5 = [[c.text.strip() for c in r.cells] for r in _t5.rows[1:]]
        if any(l and l[0] == 'Condição' for l in _linhas5):
            _ctrl = _linhas5
            break

if not _ctrl:
    erro('5: nao achei a tabela de Controle no .docx (a que carrega a Melhoria '
         '`Condição`) — ou o catalogo mudou de forma, ou a extracao parou de achar')
elif len(_ctrl) != 7:
    erro(f'5: a tabela de Controle tem {len(_ctrl)} linha(s) e eu esperava 7 — a '
         'familia mudou de tamanho e esta checagem parou de cobrir ela')
else:
    _DURA = (r'até o fim do próximo turno', r'por uma rodada', r'por 1 minuto',
             r'Dura uma rodada')
    _SAI = (r'Teste de Resistência', r'pontos de vida', r'Acaba', r'se soltar')
    _presos, _sem = [], []
    for _nome5, _custo5, _txt5 in _ctrl:
        _fala_do_alvo = re.search(r'\bO alvo\b|\bAplica uma\b', _txt5)
        _dura = any(re.search(_d, _txt5) for _d in _DURA)
        if not (_fala_do_alvo and _dura):
            continue
        _presos.append(_nome5)
        if not any(re.search(_s, _txt5) for _s in _SAI):
            _sem.append(_nome5)
    if len(_presos) < 3:
        erro(f'5: so {len(_presos)} linha(s) de Controle prendem um alvo por mais de '
             'um instante, e eu esperava pelo menos 3 (`Condição`, `Prende` e '
             '`Cerca`) — a selecao perdeu o chao e esta checagem parou de cobrir')
    elif _sem:
        erro('5: ' + ', '.join(f'`{n}`' for n in _sem) + ' prende(m) o alvo por mais '
             'de um instante e nao diz(em) como aquilo acaba — toda linha de Controle '
             'que segura alguem tem de nomear a saida, senao ela e um portao a menos '
             'pelo mesmo preco')
    else:
        print(f'  {len(_presos)} linha(s) de Controle seguram o alvo alem do instante: '
              + ', '.join(_presos))
        print('  e as ' + str(len(_presos)) + ' nomeiam a saida.')
        _fora = [n for n, _, _ in _ctrl if n not in _presos]
        print('  fora da selecao, por nao segurarem alvo nenhum: ' + ', '.join(_fora))


# --------------------------------------------------------------------------
# 6: os cinco degraus de nivel 7, comparados entre si.
#
# v0.155. O degrau do nivel 7 existe para que os cinco Caminhos recebam a mesma
# coisa, e ate aqui nada conferia isso — a regra dizia "vale exatamente o vao" e
# o vao deixou de ser um numero quando a v0.147 devolveu o ataque extra a Acao
# de Atacar. Hoje a peca 6 §3.1 publica os tres totais e declara a diferenca.
#
# Nenhum numero mora aqui: os tres saem da tabela da peca, e o teto da diferenca
# sai da frase que o declara.
# v0.158: este rotulo era `print('  6. os cinco degraus ...')`, em minuscula e
# fora do `bloco()`. O extrator da checagem 9 do conferir-repositorio.py exige
# LETRA MAIUSCULA depois do numero, entao ele nao via este bloco: a contagem do
# projeto publicava 258 e o codigo tinha 257, desde a v0.155. A guarda daquela
# checagem procura BURACO e REPETICAO, e nenhuma das duas abre aqui — o bloco
# simplesmente nao existia para ela.
bloco('6. OS CINCO DEGRAUS DE NIVEL 7 — a diferenca declarada')

_t6 = open(_p6, encoding='utf-8').read() if os.path.exists(_p6) else ''
_tot = {}
for _l in _t6.split('\n'):
    _m = re.match(r'>?\s*\|\s*\*\*(Bastião|Vanguarda)\*\*\s*\|[^|]*\|'
                  r'\s*`([\d,]+)`\s*\|\s*`([\d,]+)`\s*\|\s*\*\*`([\d,]+)`\*\*\s*\|', _l)
    if _m:
        _tot[_m.group(1)] = float(_m.group(4).replace(',', '.'))
_mg = re.search(r'Guia · Emanador · Evocador \| o degrau grande \| — \| — \| `([\d,]+)` \|', _t6)
_grande = float(_mg.group(1).replace(',', '.')) if _mg else None
_md = re.search(r'Bastião `−([\d,]+)` e Vanguarda `−([\d,]+)` contra o degrau grande', _t6)

if len(_tot) != 2 or _grande is None or not _md:
    erro('6: nao consegui ler os tres totais do nivel 7 na peca 6 §3.1 — a tabela '
         'ou a linha da diferenca declarada mudou de forma, e esta checagem parou '
         'de conferir em vez de acusar')
else:
    _db = float(_md.group(1).replace(',', '.'))
    _dv = float(_md.group(2).replace(',', '.'))
    print(f"     Bastiao {_tot['Bastião']:.2f} · Vanguarda {_tot['Vanguarda']:.2f} · "
          f"os outros tres {_grande:.2f}")
    _erro_b = abs((_grande - _tot['Bastião']) - _db)
    _erro_v = abs((_grande - _tot['Vanguarda']) - _dv)
    if _erro_b > 0.005 or _erro_v > 0.005:
        erro(f'6: a peca declara a diferenca em -{_db:.2f} e -{_dv:.2f}, e os totais '
             f'dao -{_grande-_tot["Bastião"]:.2f} e -{_grande-_tot["Vanguarda"]:.2f} '
             f'— a tabela e a frase divergiram')
    elif max(_db, _dv) > 0.50:
        erro(f'6: a diferenca declarada chegou a {max(_db,_dv):.2f} fatia, e o degrau '
             f'do nivel 7 existe para os cinco receberem a mesma coisa — acima de '
             f'0,50 ela deixa de ser residuo e vira degrau desigual')
    else:
        print('     [x] os totais batem com a diferenca declarada, e ela cabe no teto.')


# --------------------------------------------------------------------------
print()
print('=' * 88)
if FALHAS:
    print(f'>>> {len(FALHAS)} PROBLEMA(S):')
    for e in FALHAS:
        print('   -', e)
    sys.exit(1)
print('>>> TUDO OK — nenhum vocabulario de outro sistema no manual, nenhum termo')
print('    mecanico indefinido sem motivo, e os numeros importados continuam batendo.')
if AVISOS:
    print(f'    {len(AVISOS)} aviso(s) acima, que nao falham o validador.')
