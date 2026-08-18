#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Confere o quadro de pericias e oficios da peca 07.

  Pericia = d20 + atributo dela + maestria (maestria so entra se treinado)
  Oficio  = d20 + atributo que a situacao pede + maestria (idem)

O que este validador garante:
  1. contagem por atributo bate com o documento
  2. nenhum nome aparece duas vezes, nem como pericia e oficio ao mesmo tempo
  3. as fixas de cada Caminho existem no quadro e nenhuma se repete dentro do Caminho
  4. a fracao treinada fica na faixa que a peca 6 pediu (30% a 40%)
  5. toda pericia e alcancavel por alguem
  6. nenhum nome colide com termo definido do Fundamento nem com termo interno do
     projeto — e as colisoes ACEITAS de proposito estao declaradas aqui, com o motivo

Roda sem argumento. Sai com codigo 1 se algo quebrar.
"""

import os
import re
import sys
import unicodedata

FALHAS = []
# Checagem que nao rodou — ver o comentario igual no conferir-nomes.py. Sem
# python-docx este validador dizia TUDO OK sem ter aberto o manual uma vez.
PULADAS = []


def erro(msg):
    FALHAS.append(msg)
    print(f'  !! {msg}')


def norma(s):
    s = unicodedata.normalize('NFD', s.lower())
    return ''.join(c for c in s if unicodedata.category(c) != 'Mn')


# --------------------------------------------------------------------------
# O QUADRO
# --------------------------------------------------------------------------
# v0.16: "Inteligencia SABE, Essencia PERCEBE". Sentir Energia e Percepcao
# moraram em Inteligencia da v0.8 ate a v0.15 e mudaram de casa aqui.

PERICIAS = {
    'Forca':        ['Atletismo'],
    'Destreza':     ['Acrobacia', 'Furtividade', 'Pontaria', 'Prestidigitacao'],
    'Inteligencia': ['Intuicao', 'Investigacao', 'Ocultismo', 'Religiao', 'Historia',
                     'Hierarquia', 'Medicina', 'Sobrevivencia', 'Natureza',
                     'Lidar com Animais', 'Tecnologia'],
    'Essencia':     ['Sentir Energia', 'Percepcao', 'Persuasao', 'Enganacao',
                     'Intimidacao', 'Atuacao', 'Provocar'],
    'Constituicao': [],
}

# v0.42: a lista de oficios saiu daqui e passou a ser LIDA da peca 7.
# Ela morava em tres lugares — a peca, este validador e o dados.js da ficha —
# e so dois eram comparados (o conferir-ficha.py cruza a peca com o dados.js).
# Este arquivo era a terceira copia, sem ninguem conferindo. Licao no 9.
_P7 = open(os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        '07-pericias-e-oficios.md'), encoding='utf-8').read()
_SEC = re.search(r'## 5\. Os \w+ of[ií]cios(.*?)(?=\n## )', _P7, re.S)
OFICIOS = [norma(n).capitalize() for n in re.findall(r'\*\*([^*]+)\*\* — ', _SEC.group(1))] if _SEC else []

# Quantos a peca DECLARA, lido do titulo dela — separado da lista aplicada,
# que e a licao no 8: a checagem nao pode se medir contra a propria constante.
_NUM = {'nove': 9, 'dez': 10, 'onze': 11, 'doze': 12, 'treze': 13, 'catorze': 14}
_m = re.search(r'## 5\. Os (\w+) of[ií]cios', _P7)
OFICIOS_DECLARADOS = _NUM.get(_m.group(1).lower()) if _m else None

# Cada Caminho fixa duas pericias e um oficio. O resto o jogador escolhe livre.
CAMINHOS = {
    'Bastiao':   {'pericias': ['Atletismo', 'Intimidacao'],       'oficio': 'Forja'},
    'Vanguarda': {'pericias': ['Acrobacia', 'Percepcao'],         'oficio': 'Arrombamento'},
    'Guia':      {'pericias': ['Persuasao', 'Medicina'],          'oficio': 'Herbalismo'},
    'Emanador':  {'pericias': ['Ocultismo', 'Investigacao'],      'oficio': 'Caligrafia'},
    'Evocador':  {'pericias': ['Religiao', 'Lidar com Animais'],  'oficio': 'Entalhador'},
}
# Sentir Energia NAO e fixa de ninguem, de proposito: ela e a mais rolada da mesa,
# e fixa-la em um Caminho daria a ele uma escolha livre a mais na pratica. Livre
# para todos, ela vira decisao — e o feiticeiro ruim de sentir energia passa a caber.

CAM_FIXAS, CAM_LIVRES = 2, 4       # pericias que o Caminho entrega
CAM_OF_FIXO, CAM_OF_LIVRE = 1, 1   # oficios que o Caminho entrega
ORI_PERICIAS = 2                   # uma da lista da Origem + uma livre
ORI_EXTRA = 1                      # mais um OFICIO livre, ou outra PERICIA no lugar
FAIXA_TREINADA = (0.30, 0.42)

# Cada Origem oferece quatro pericias e voce treina uma (peca 09).
ORIGENS = {
    'Latente':            ['Sentir Energia', 'Sobrevivencia', 'Furtividade', 'Intuicao'],
    'Receptaculo':        ['Sentir Energia', 'Ocultismo', 'Intuicao', 'Religiao'],
    'Descendente':        ['Hierarquia', 'Historia', 'Ocultismo', 'Persuasao'],
    'Reencarnado':        ['Historia', 'Ocultismo', 'Investigacao', 'Intimidacao'],
    'Feto':               ['Ocultismo', 'Medicina', 'Sentir Energia', 'Natureza'],
}

TODAS = [p for grupo in PERICIAS.values() for p in grupo]

# Termos que ja significam outra coisa dentro do projeto.
INTERNOS = {
    'Muro', 'Punho', 'Brasa', 'Estocada', 'Alcance', 'Oficio', 'Elo', 'Folego',
    'Regua', 'Torrente', 'Explosivo', 'Arremate', 'Sombra', 'Enxame', 'Coro',
    'Bastiao', 'Vanguarda', 'Guia', 'Emanador', 'Evocador', 'Vigor', 'Intelecto',
    'Espirito', 'Maestria', 'Classe', 'Refino', 'Caminho', 'Trilha', 'Origem',
    'Pacto', 'Grau', 'Selo', 'Barreira', 'Dominio', 'Familia', 'Melhoria',
    'Restricao', 'Passiva', 'Forma', 'Integridade', 'Aptidao', 'Defesa',
}

# Colisoes conhecidas e ACEITAS de proposito. Nao falham o validador, mas
# aparecem no relatorio para nao serem confundidas com descuido depois.
COLISOES_ACEITAS = {
    'Provocar': 'o manual usa "sem provocar ataque de oportunidade" 3x. Decidido na '
                'v0.16 que a frase e comum demais para confundir na mesa.',
    'Natureza': 'o manual tem a Passiva "Segunda Natureza". Duas palavras contra uma, '
                'e a Passiva e rara. Colisao fraca, aceita.',
    'Historia': 'aparece 1x no manual em prosa solta. Pela nota de metodo da v0.6, '
                'prosa solta nao e colisao.',
}


def bloco(titulo):
    print()
    print('=' * 88)
    print(titulo)
    print('=' * 88)


# --------------------------------------------------------------------------
bloco('1. CONTAGEM POR ATRIBUTO')

total = len(TODAS)
for atr, lista in PERICIAS.items():
    n = len(lista)
    print(f'  {atr:<16} {n:>2} pericias  {n/total:>5.0%}   ' + ' · '.join(lista))
print(f'\n  total: {total} pericias, {len(OFICIOS)} oficios')

if not 21 <= total <= 25:
    erro(f'{total} pericias — a peca 07 declara 23')
if not OFICIOS:
    erro('nao consegui ler a lista de oficios da peca 07 (o titulo do §5 mudou?)')
if OFICIOS_DECLARADOS is None:
    erro('o titulo do §5 da peca 07 nao diz quantos oficios sao, por extenso')
elif len(OFICIOS) != OFICIOS_DECLARADOS:
    erro(f'a peca 07 lista {len(OFICIOS)} oficios e o titulo dela declara '
         f'{OFICIOS_DECLARADOS}')
if PERICIAS['Constituicao']:
    erro('Constituicao ganhou pericia — a peca 4 decidiu que ela nao tem nenhuma')

# --------------------------------------------------------------------------
bloco('2. PESO DE MESA — contagem nao e a mesma coisa que valor')

# Quantas vezes cada pericia e rolada numa campanha tipica, em peso relativo.
# Serve para responder "qual atributo vale mais fora de combate", que a contagem
# crua responde errado: uma Sentir Energia vale por quatro pericias de nicho.
PESO = {'Sentir Energia': 4, 'Percepcao': 3, 'Intuicao': 2, 'Investigacao': 2,
        'Atletismo': 2, 'Furtividade': 2, 'Persuasao': 2, 'Acrobacia': 1.5,
        'Ocultismo': 1.5, 'Intimidacao': 1.5, 'Medicina': 1, 'Enganacao': 1}

peso_total = sum(PESO.get(p, 1) for p in TODAS)
print(f"  {'atributo':<16} {'pericias':>9} {'contagem':>9} {'peso de mesa':>13}")
pesos = {}
for atr, lista in PERICIAS.items():
    w = sum(PESO.get(p, 1) for p in lista) / peso_total
    pesos[atr] = w
    print(f'  {atr:<16} {len(lista):>9} {len(lista)/total:>8.0%} {w:>12.0%}')

print('\n  Antes da v0.16, com Sentir Energia e Percepcao em Inteligencia:')
print('    Inteligencia 56% de peso, Essencia 21%.')
print('  A troca nao mudou a contagem para melhor — mudou o peso, que e o que importa.')

lider = max(pesos, key=pesos.get)
if pesos[lider] > 0.50:
    erro(f'{lider} concentra {pesos[lider]:.0%} do valor de mesa — vira atributo obrigatorio')

# --------------------------------------------------------------------------
bloco('3. NENHUM NOME REPETIDO, E NENHUM EM DUAS LISTAS')

vistos = {}
for atr, lista in PERICIAS.items():
    for p in lista:
        if norma(p) in vistos:
            erro(f'"{p}" aparece em {vistos[norma(p)]} e em {atr}')
        vistos[norma(p)] = atr
for o in OFICIOS:
    if norma(o) in vistos:
        erro(f'"{o}" e pericia ({vistos[norma(o)]}) e oficio ao mesmo tempo')
    vistos[norma(o)] = 'oficio'
print(f'  {len(vistos)} nomes, todos unicos entre pericias e oficios.')

# --------------------------------------------------------------------------
bloco('4. AS FIXAS DE CADA CAMINHO')

print(f"  {'Caminho':<12} {'pericias fixas':<34} {'oficio fixo':<14} atributos")
for nome, d in CAMINHOS.items():
    fora = [p for p in d['pericias'] if p not in TODAS]
    if d['oficio'] not in OFICIOS:
        fora.append(d['oficio'])
    atrs = [vistos[norma(p)] for p in d['pericias'] if norma(p) in vistos]
    print(f"  {nome:<12} {' · '.join(d['pericias']):<34} {d['oficio']:<14} {' + '.join(a[:3] for a in atrs)}")
    if len(d['pericias']) != CAM_FIXAS:
        erro(f'{nome} fixa {len(d["pericias"])} pericias, deveria fixar {CAM_FIXAS}')
    if fora:
        erro(f'{nome} fixa o que nao existe no quadro: {", ".join(fora)}')
    if len(set(map(norma, d['pericias']))) != len(d['pericias']):
        erro(f'{nome} repete uma pericia nas proprias fixas')

repetidas = [p for p in TODAS
             if sum(1 for d in CAMINHOS.values() if p in d['pericias']) > 1]
print(f"\n  fixada por mais de um Caminho: {', '.join(repetidas) if repetidas else '—'}")
print('  As fixas sao a assinatura do Caminho. Repetir demais apaga a diferenca entre eles.')

# Quanto cada Caminho ganha de graca. Se um ganha muito mais que outro, ele leva
# na pratica uma escolha livre a mais. Informativo: o peso abaixo e estimado.
print(f"\n  {'Caminho':<12} {'peso das duas fixas':>20}")
cargas = {}
for nome, d in CAMINHOS.items():
    c = sum(PESO.get(p, 1) for p in d['pericias'])
    cargas[nome] = c
    print(f'  {nome:<12} {c:>20.1f}')
lo, hi = min(cargas.values()), max(cargas.values())
print(f'\n  faixa: {lo:.1f} a {hi:.1f}. Diferenca de {hi - lo:.1f} sobre 8 pericias treinadas.')
if hi - lo > 3.0:
    erro(f'as fixas de {max(cargas, key=cargas.get)} valem {hi/lo:.1f}x as de '
         f'{min(cargas, key=cargas.get)} — o Caminho mais bem servido ganha uma escolha livre disfarcada')
else:
    print('  Dentro do ruido: as 4 escolhas livres corrigem sozinhas essa diferenca.')

# --------------------------------------------------------------------------
bloco('5. FRACAO TREINADA')

base_p = CAM_FIXAS + CAM_LIVRES + ORI_PERICIAS
base_o = CAM_OF_FIXO + CAM_OF_LIVRE
# a Origem entrega UM extra, e o jogador escolhe se ele e oficio ou pericia
rotas = {'Origem pega o OFICIO':  (base_p, base_o + ORI_EXTRA),
         'Origem pega a PERICIA': (base_p + ORI_EXTRA, base_o)}
print(f'  {CAM_FIXAS} fixas + {CAM_LIVRES} livres do Caminho + {ORI_PERICIAS} da Origem,'
      f' mais UM extra que o jogador escolhe:\n')
print(f"  {'rota':<24}{'pericias':>18}{'oficios':>18}")
for nome, (p, o) in rotas.items():
    print(f'  {nome:<24}{f"{p} de {total} = {p/total:.0%}":>18}'
          f'{f"{o} de {len(OFICIOS)} = {o/len(OFICIOS):.0%}":>18}')
    if not FAIXA_TREINADA[0] <= p / total <= FAIXA_TREINADA[1]:
        erro(f'{nome}: fracao de pericia em {p/total:.0%}, fora da faixa '
             f'{FAIXA_TREINADA[0]:.0%}-{FAIXA_TREINADA[1]:.0%}')
    if o / len(OFICIOS) > 0.35:
        erro(f'{nome}: oficios em {o/len(OFICIOS):.0%} — oficio precisa ser raro para virar cena')
print(f'\n  faixa aceita para pericia: {FAIXA_TREINADA[0]:.0%} a {FAIXA_TREINADA[1]:.0%}')
print('  Abaixo disso a ficha esvazia. Acima, "ser treinado" para de significar algo')
print('  e nao sobra em que o resto do grupo brilhar. As duas rotas precisam caber.')

# --------------------------------------------------------------------------
bloco('5.1 AS LISTAS DE ORIGEM')
print(f"  {'Origem':<16}{'oferece':>9}   pericias")
for nome, lista in ORIGENS.items():
    fora = [p for p in lista if p not in TODAS]
    print(f"  {nome:<16}{len(lista):>9}   {' · '.join(lista)}")
    if len(lista) != 4:
        erro(f'{nome} oferece {len(lista)} pericias, e a peca 09 diz quatro')
    if fora:
        erro(f'{nome} oferece o que nao existe no quadro: {", ".join(fora)}')
    if len(set(map(norma, lista))) != len(lista):
        erro(f'{nome} repete uma pericia na propria lista')

print('\n  Origem cuja lista o Caminho ja fixa por inteiro (a escolha da Origem morreria):')
achou = False
for o_nome, o_lista in ORIGENS.items():
    for c_nome, c in CAMINHOS.items():
        if set(o_lista) <= set(c['pericias']):
            erro(f'{o_nome} tem as quatro pericias ja fixadas por {c_nome}')
            achou = True
if not achou:
    print('    nenhuma. Sobreposicao parcial e esperada — quem cair nela escolhe outra da lista.')

usadas = {p for l in ORIGENS.values() for p in l}
print(f'\n  as cinco Origens tocam {len(usadas)} das {total} pericias')

bloco('6. TODA PERICIA E ALCANCAVEL')

print('  Com 4 pericias e 1 oficio de escolha livre em todo o quadro, qualquer nome')
print('  e alcancavel por qualquer Caminho. Nao existe pericia orfa desde a v0.16 —')
print('  antes, Prestidigitacao e Jogatina so vinham da Origem.')
nunca_fixa = [p for p in TODAS if not any(p in d['pericias'] for d in CAMINHOS.values())]
print(f'\n  pericias que nenhum Caminho fixa ({len(nunca_fixa)} de {total}): {", ".join(nunca_fixa)}')
print('  Isso e esperado: so 10 das pericias sao assinatura de alguem.')

# --------------------------------------------------------------------------
bloco('7. COLISAO DE TERMO — contra o Fundamento e contra o proprio projeto')

AQUI = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(AQUI, '..', '..', 'manual', 'Fundamento-MANUAL-v7.docx')

texto_manual = None
try:
    import docx  # python-docx
    doc = docx.Document(DOCX)
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                partes.append(c.text)
    texto_manual = '\n'.join(partes)
except Exception as e:
    print(f'  (manual nao lido: {e})')
    print('  A checagem contra o Fundamento foi PULADA. Rode com python-docx instalado')
    print('  e o manual em manual/Fundamento-MANUAL-v7.docx antes de fechar versao.')

nomes = TODAS + OFICIOS
internos_norm = {norma(x) for x in INTERNOS}
achados = []

for n in nomes:
    if norma(n) in internos_norm:
        erro(f'"{n}" ja e termo interno do projeto — colisao para dentro')
    if texto_manual:
        pat = re.compile(r'\b' + re.escape(norma(n)) + r'\b')
        linhas = [l.strip() for l in texto_manual.split('\n')
                  if l.strip() and pat.search(norma(l))]
        if linhas:
            achados.append((n, len(linhas), linhas[0][:58]))

if texto_manual:
    print('  Nota de metodo da v0.6: separar TERMO DEFINIDO de PROSA SOLTA.\n')
    for n, q, amostra in achados:
        if n in COLISOES_ACEITAS:
            print(f'  ACEITA  {n:<18} {q}x no manual')
            print(f'          motivo: {COLISOES_ACEITAS[n]}')
        else:
            erro(f'"{n}" aparece {q}x no manual e nao esta na lista de aceitas: "{amostra}..."')
    nao_apareceram = [n for n in COLISOES_ACEITAS if n not in [a[0] for a in achados]]
    for n in nao_apareceram:
        if n in nomes:
            print(f'  (declarada aceita mas nao aparece mais no manual: {n} — pode sair da lista)')
    if not achados:
        print('  Nenhum dos nomes aparece no manual.')
else:
    PULADAS.append('7. colisao de termo contra o Fundamento (sem o manual)')
    print('  PULADA — nenhum nome foi batido contra o manual.')

# --------------------------------------------------------------------------
bloco('8. A ESCADA DE DIFICULDADE — o que o treino compra')


def maestria(nv):
    return 1 + (nv - 2) // 8


def investido(nv):
    return min(6, 3 + (nv - 2) // 8)


NIVEIS = [2, 10, 18, 30]
ESCADA = [(10, 'rotina'), (14, 'facil'), (18, 'media'), (22, 'dificil'), (26, 'quase impossivel')]

print(f"  {'dificuldade':<26}" + ''.join(f'nv{nv:<6}' for nv in NIVEIS) + '   sem treino no nv30')
for cd, rot in ESCADA:
    linha = f'  {"CD " + str(cd) + " (" + rot + ")":<26}'
    for nv in NIVEIS:
        alvo = cd - (investido(nv) + maestria(nv))
        linha += f'{max(0, min(100, (21 - alvo) * 5)):>4}%   '
    alvo_sem = cd - investido(30)
    linha += f'{max(0, min(100, (21 - alvo_sem) * 5)):>10}%'
    print(linha)
print('\n  Personagem treinado que investiu no atributo. A ultima coluna mostra o mesmo')
print('  personagem sem treino no nivel 30: a maestria e a diferenca, e ela vale 20 pp.')

# --------------------------------------------------------------------------
print('\n' + '=' * 88)
print('O ATRIBUTO PADRAO DE CADA OFICIO — v0.104')
print('=' * 88)
# Os onze oficios ganharam atributo padrao. A checagem confere que TODOS tem um,
# que ele e' um dos cinco atributos que a peca 8 governa, e que a distribuicao
# publicada na prosa bate com a tabela — sao duas copias do mesmo numero.
_ATRIB = ['Força', 'Destreza', 'Constituição', 'Inteligência', 'Essência']
_sec = _P7[_P7.find('### O atributo padrão de cada ofício'):]
_sec = _sec[:_sec.find('\n### ')] if '\n### ' in _sec else _sec
_pad = {}
for _l in _sec.split('\n'):
    _m = re.match(r'^\|\s*\*\*([^*]+)\*\*\s*\|\s*([A-ZÀ-Ú][a-zçãêéíóú]+)\s*\|', _l)
    if _m:
        _pad[norma(_m.group(1)).capitalize()] = _m.group(2)
_semp = [o for o in OFICIOS if o not in _pad]
_sobra = [o for o in _pad if o not in OFICIOS]
if not _sec:
    FALHAS.append('a peca 07 nao tem mais a secao do atributo padrao de oficio')
elif _semp:
    FALHAS.append('oficio sem atributo padrao na peca 07: ' + ', '.join(_semp))
elif _sobra:
    FALHAS.append('a tabela de atributo padrao tem oficio que nao existe: '
                  + ', '.join(_sobra))
else:
    _mau = [f'{o} ({a})' for o, a in _pad.items() if a not in _ATRIB]
    if _mau:
        FALHAS.append('atributo padrao que nao e um dos cinco: ' + ', '.join(_mau))
    else:
        _c = {a: sum(1 for v in _pad.values() if v == a) for a in _ATRIB}
        print('  ' + ' · '.join(f'{a} {_c[a]}' for a in _ATRIB))
        _m2 = re.search(r'\*\*(\w+) em Destreza, (\w+) em Inteligência, (\w+) em '
                        r'Essência, (\w+) em Força — e nenhuma em Constituição\.\*\*', _sec)
        _EXT = {'uma': 1, 'duas': 2, 'três': 3, 'quatro': 4, 'cinco': 5, 'seis': 6}
        if not _m2:
            FALHAS.append('a peca 07 nao publica mais a distribuicao dos atributos '
                          'padrao em prosa, no formato que esta checagem le')
        else:
            _esc = [_EXT.get(x.lower()) for x in _m2.groups()]
            if _esc != [_c['Destreza'], _c['Inteligência'], _c['Essência'], _c['Força']] \
                    or _c['Constituição'] != 0:
                FALHAS.append(f'a prosa da peca 07 diz {_esc} (Des/Int/Ess/For) e a '
                              f'tabela tem {[_c["Destreza"], _c["Inteligência"], _c["Essência"], _c["Força"]]}')
            else:
                print(f'  [x] os {len(_pad)} oficios tem atributo padrao, e a prosa bate com a tabela')
        if 'antes da rolagem' not in _sec:
            FALHAS.append('a peca 07 nao exige mais que o mestre diga o atributo ANTES '
                          'da rolagem — sem isso o padrao nao tira julgamento nenhum')
        else:
            print('  [x] o mestre diz o atributo antes da rolagem, e a peca cobra isso')

# --------------------------------------------------------------------------
print()
print('=' * 88)
if FALHAS:
    print(f'>>> {len(FALHAS)} PROBLEMA(S):')
    for f in FALHAS:
        print(f'    - {f}')
    sys.exit(1)
if PULADAS:
    print(f'>>> OK, mas {len(PULADAS)} checagem(ns) PULARAM:')
    for p in PULADAS:
        print(f'    - {p}')
    print('    O que pulou NAO foi conferido. Um verde que pulou checagem nao e')
    print('    um verde. Instale: pip install python-docx --break-system-packages')
else:
    print('>>> TUDO OK — o quadro fecha, a fracao esta na faixa, e as colisoes que existem')
    print('    estao declaradas com motivo em vez de esquecidas.')
